"""
Ingestion Service - Document processing and chunking
"""

import asyncio
import io
import uuid
import json
import math
import hashlib
import re
import zipfile
from typing import Any, AsyncIterator, List, Optional, Tuple
from datetime import datetime, timedelta, timezone
from fastapi import UploadFile
import structlog
import httpx

from commons.types.requests import IngestionRequest
from commons.types.responses import IngestionResponse, IngestionStatus
from ..config import Settings
from ..connections import connection_manager
from .chunking import (
    DEFAULT_CHUNK_OVERLAP,
    DEFAULT_CHUNK_SIZE,
    clamp_chunking,
    chunk_text,
    resolve_agent_chunking,
)
from .job_store import (
    DURABLE_REINDEX_KIND,
    DuplicateDurableJobError,
    JobStore,
    JobStoreUnavailableError,
)
from .ingestion_payload_store import (
    IngestionPayloadStore,
    InvalidIngestionPayloadError,
)
from ..security.malware_scanner import (
    MalwareDetectedError,
    MalwareScanner,
    MalwareScannerUnavailableError,
)
from .qdrant_vector_service import QdrantVectorService
from .runtime_settings_service import RuntimeSettingsService

logger = structlog.get_logger()

STAGED_CHUNKS_COLLECTION = "ingestion_staged_chunks"
_IDEMPOTENCY_KEY_PATTERN = re.compile(r"^[A-Za-z0-9][A-Za-z0-9._:-]{0,127}$")


class IngestionEmbeddingError(RuntimeError):
    """Raised when Voyage cannot produce a valid embedding."""


class IngestionStorageError(RuntimeError):
    """Raised when a chunk cannot be durably stored in every configured backend."""


class InvalidIngestionInputError(ValueError):
    """A permanent, safe-to-report source/input validation failure."""


class IngestionIdempotencyConflictError(ValueError):
    """An idempotency key was reused with a different upload manifest."""


def _utc_timestamp() -> str:
    """Return an unambiguous UTC timestamp for durable job state."""
    return datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")


def _public_job_error(exc: Exception) -> str:
    """Map backend failures to a safe status message for dashboard clients."""
    if isinstance(exc, IngestionEmbeddingError):
        return "Document embedding failed"
    if isinstance(exc, IngestionStorageError):
        return "Failed to store knowledge-base document"
    if isinstance(exc, (InvalidIngestionInputError, InvalidIngestionPayloadError)):
        return "Document input is invalid"
    return "Document processing failed"


_SAFE_JOB_ERRORS = {
    "Document embedding failed",
    "Failed to store knowledge-base document",
    "Document processing failed",
    "Document input is invalid",
}


def _safe_stored_job_error(status: str, error: object) -> Optional[str]:
    """Keep legacy job records from exposing stored provider diagnostics."""
    if status != "error":
        return None
    if isinstance(error, str) and error in _SAFE_JOB_ERRORS:
        return error
    return "Document processing failed"


class IngestionService:
    """Service for document ingestion and processing."""
    
    def __init__(self, settings: Settings):
        self.settings = settings
        self.job_store = JobStore(job_ttl_seconds=settings.INGESTION_JOB_TTL_SECONDS)
        self.payload_store = IngestionPayloadStore(settings)
        self.malware_scanner = MalwareScanner(settings)
        self.runtime_settings_service = RuntimeSettingsService(settings)
        self.qdrant = QdrantVectorService(settings) if settings.VECTOR_BACKEND == "qdrant" else None

    async def _get_voyage_runtime_config(self) -> dict[str, str]:
        config = await self.runtime_settings_service.get_voyage_runtime_config()
        return {
            "api_key": config["api_key"],
            "base_url": config["base_url"],
            "model": config["model"],
        }
    
    async def start_ingestion_job(
        self,
        files: List[dict],
        agent_id: Optional[str] = None,
        brand_id: Optional[str] = None,
    ) -> str:
        """Start a new ingestion job."""
        job_id = str(uuid.uuid4())

        await self.job_store.set(job_id, {
            "status": "pending",
            "files_count": len(files),
            "processed_count": 0,
            "agent_id": agent_id,
            "brand_id": brand_id,
            "created_at": _utc_timestamp(),
            "completed_at": None,
            "error": None,
        })

        logger.info("Started ingestion job", job_id=job_id, files_count=len(files), agent_id=agent_id)
        return job_id

    async def submit_durable_job(
        self,
        files: List[dict],
        *,
        agent_id: str | None,
        brand_id: str,
        brand_slug: str,
        chunk_size: int,
        chunk_overlap: int,
        idempotency_key: str | None = None,
        job_metadata: dict[str, Any] | None = None,
    ) -> str:
        """Persist encrypted sources and an immutable v2 job before returning.

        This is deliberately separate from ``start_ingestion_job`` because the
        latter remains a compatibility hook for historical tests/callers. Only
        this method creates work that the restart-safe worker is allowed to
        claim.
        """
        if not all(isinstance(value, str) and value for value in (brand_id, brand_slug)):
            raise InvalidIngestionInputError("Agent brand scope is incomplete")
        if agent_id is not None and (not isinstance(agent_id, str) or not agent_id):
            raise InvalidIngestionInputError("Agent identifier is invalid")
        if not files:
            raise InvalidIngestionInputError("At least one document is required")
        normalized_idempotency_key = self._normalize_idempotency_key(idempotency_key)
        normalized_job_metadata = self._normalize_durable_job_metadata(job_metadata)
        source_manifest = self._source_manifest(files)
        if normalized_idempotency_key:
            existing = await self.job_store.find_durable_job_by_idempotency(
                agent_id=agent_id,
                brand_id=brand_id,
                idempotency_key=normalized_idempotency_key,
            )
            if existing:
                if self._matches_idempotent_submission(
                    existing,
                    source_manifest=source_manifest,
                    brand_slug=brand_slug,
                ):
                    return str(existing["job_id"])
                raise IngestionIdempotencyConflictError(
                    "Idempotency key was already used for a different document upload"
                )
        # Source bytes must be scanned before durable encryption/persistence so
        # an infected payload cannot be retained, leased, or processed later.
        for file_index, file_data in enumerate(files):
            try:
                scan = await self.malware_scanner.scan(
                    file_data["content"], filename=file_data["filename"]
                )
                logger.info(
                    "durable_upload_malware_scan_completed",
                    brand_id=brand_id,
                    file_index=file_index,
                    result=scan.status,
                )
            except MalwareDetectedError:
                logger.warning(
                    "durable_upload_malware_detected",
                    brand_id=brand_id,
                    file_index=file_index,
                )
                raise
            except MalwareScannerUnavailableError:
                logger.error(
                    "durable_upload_malware_scan_unavailable",
                    brand_id=brand_id,
                    file_index=file_index,
                )
                raise
        job_id = str(uuid.uuid4())
        snapshot = {
            "agent_id": agent_id,
            "brand_id": brand_id,
            "brand_slug": brand_slug,
            "chunk_size": int(chunk_size),
            "chunk_overlap": int(chunk_overlap),
        }
        payload_refs = await self.payload_store.persist(job_id, files)
        try:
            await self.job_store.create_durable_job(
                job_id,
                {
                    "files_count": len(payload_refs),
                    "agent_id": agent_id,
                    "brand_id": brand_id,
                    "brand_slug": brand_slug,
                    "snapshot": snapshot,
                    "payload_refs": payload_refs,
                    "source_manifest": source_manifest,
                    **normalized_job_metadata,
                    **({"idempotency_key": normalized_idempotency_key} if normalized_idempotency_key else {}),
                    "max_attempts": max(1, int(self.settings.INGESTION_MAX_ATTEMPTS)),
                },
            )
        except DuplicateDurableJobError:
            await self.payload_store.delete_job_payloads(job_id)
            existing = await self.job_store.find_durable_job_by_idempotency(
                agent_id=agent_id,
                brand_id=brand_id,
                idempotency_key=normalized_idempotency_key or "",
            )
            if existing and self._matches_idempotent_submission(
                existing,
                source_manifest=source_manifest,
                brand_slug=brand_slug,
            ):
                return str(existing["job_id"])
            raise IngestionIdempotencyConflictError(
                "Idempotency key was already used for a different document upload"
            )
        except Exception:
            # A payload without its job is inaccessible and is promptly removed;
            # the TTL index remains a bounded-retention backstop on interruption.
            await self.payload_store.delete_job_payloads(job_id)
            raise
        logger.info("durable_ingestion_job_queued", job_id=job_id, files_count=len(payload_refs))
        return job_id

    async def submit_reindex_job(
        self,
        *,
        document_id: str,
        agent_id: str | None,
        brand_id: str,
        brand_slug: str,
        chunk_size: int,
        chunk_overlap: int,
        idempotency_key: str | None = None,
    ) -> str:
        """Queue a durable in-place embedding refresh for one knowledge document.

        A re-index refreshes embeddings and vector-backend payloads from already
        published chunks; it does not need to retain source bytes or create a
        second logical document. The immutable target and tenant snapshot make
        retries, cancellation, and worker recovery use the same protocol as
        uploads.
        """
        if not all(isinstance(value, str) and value for value in (document_id, brand_id, brand_slug)):
            raise InvalidIngestionInputError("Knowledge document scope is incomplete")
        if agent_id is not None and (not isinstance(agent_id, str) or not agent_id):
            raise InvalidIngestionInputError("Agent identifier is invalid")
        normalized_idempotency_key = self._normalize_idempotency_key(idempotency_key)
        if normalized_idempotency_key:
            existing = await self.job_store.find_durable_job_by_idempotency(
                agent_id=agent_id,
                brand_id=brand_id,
                idempotency_key=normalized_idempotency_key,
                kind=DURABLE_REINDEX_KIND,
            )
            if existing:
                if existing.get("reindex_document_id") == document_id and existing.get("brand_slug") == brand_slug:
                    return str(existing["job_id"])
                raise IngestionIdempotencyConflictError(
                    "Idempotency key was already used for a different re-index request"
                )

        job_id = str(uuid.uuid4())
        snapshot = {
            "agent_id": agent_id,
            "brand_id": brand_id,
            "brand_slug": brand_slug,
            "chunk_size": int(chunk_size),
            "chunk_overlap": int(chunk_overlap),
        }
        try:
            await self.job_store.create_durable_job(
                job_id,
                {
                    "files_count": 0,
                    "items_count": 1,
                    "agent_id": agent_id,
                    "brand_id": brand_id,
                    "brand_slug": brand_slug,
                    "snapshot": snapshot,
                    "reindex_document_id": document_id,
                    "submission_kind": "knowledge_reindex",
                    **({"idempotency_key": normalized_idempotency_key} if normalized_idempotency_key else {}),
                    "max_attempts": max(1, int(self.settings.INGESTION_MAX_ATTEMPTS)),
                },
                kind=DURABLE_REINDEX_KIND,
            )
        except DuplicateDurableJobError:
            existing = await self.job_store.find_durable_job_by_idempotency(
                agent_id=agent_id,
                brand_id=brand_id,
                idempotency_key=normalized_idempotency_key or "",
                kind=DURABLE_REINDEX_KIND,
            )
            if existing and existing.get("reindex_document_id") == document_id and existing.get("brand_slug") == brand_slug:
                return str(existing["job_id"])
            raise IngestionIdempotencyConflictError(
                "Idempotency key was already used for a different re-index request"
            )
        logger.info("durable_knowledge_reindex_queued", job_id=job_id, brand_id=brand_id)
        return job_id

    @staticmethod
    def _normalize_idempotency_key(value: str | None) -> str | None:
        if value is None:
            return None
        if not isinstance(value, str) or not _IDEMPOTENCY_KEY_PATTERN.fullmatch(value):
            raise InvalidIngestionInputError("Idempotency key is invalid")
        return value

    @staticmethod
    def _normalize_durable_job_metadata(value: dict[str, Any] | None) -> dict[str, Any]:
        """Allow only small, source-free status metadata in a durable job."""
        if value is None:
            return {}
        if not isinstance(value, dict):
            raise InvalidIngestionInputError("Job metadata is invalid")
        allowed = {"submission_kind", "items_count"}
        if set(value) - allowed:
            raise InvalidIngestionInputError("Job metadata contains unsupported fields")
        normalized: dict[str, Any] = {}
        if "submission_kind" in value:
            kind = value["submission_kind"]
            if not isinstance(kind, str) or kind not in {"document", "knowledge_document", "knowledge_bulk", "knowledge_reindex"}:
                raise InvalidIngestionInputError("Job submission kind is invalid")
            normalized["submission_kind"] = kind
        if "items_count" in value:
            try:
                items_count = int(value["items_count"])
            except (TypeError, ValueError) as exc:
                raise InvalidIngestionInputError("Job item count is invalid") from exc
            if not 1 <= items_count <= 1000:
                raise InvalidIngestionInputError("Job item count is invalid")
            normalized["items_count"] = items_count
        return normalized

    @staticmethod
    def _source_manifest(files: List[dict]) -> list[dict[str, Any]]:
        manifest: list[dict[str, Any]] = []
        for index, file_data in enumerate(files):
            content = file_data.get("content")
            filename = file_data.get("filename")
            content_type = file_data.get("content_type")
            context = file_data.get("context")
            if not isinstance(content, bytes) or not isinstance(filename, str) or not isinstance(content_type, str):
                raise InvalidIngestionInputError("Document upload is invalid")
            if context is not None:
                try:
                    context_serialized = json.dumps(context, sort_keys=True, separators=(",", ":"))
                except (TypeError, ValueError) as exc:
                    raise InvalidIngestionInputError("Document upload context is invalid") from exc
                context_sha256 = hashlib.sha256(context_serialized.encode("utf-8")).hexdigest()
            else:
                context_sha256 = None
            manifest.append(
                {
                    "file_index": index,
                    "filename": filename,
                    "content_type": content_type,
                    "size_bytes": len(content),
                    "sha256": hashlib.sha256(content).hexdigest(),
                    **({"context_sha256": context_sha256} if context_sha256 else {}),
                }
            )
        return manifest

    @staticmethod
    def _matches_idempotent_submission(
        existing: dict,
        *,
        source_manifest: list[dict[str, Any]],
        brand_slug: str,
    ) -> bool:
        return (
            existing.get("brand_slug") == brand_slug
            and existing.get("source_manifest") == source_manifest
        )

    async def process_next_durable_job(self, *, worker_id: str) -> bool:
        """Claim and process at most one v2 job. Used by the standalone worker."""
        job = await self.job_store.claim_next(
            worker_id=worker_id,
            lease_seconds=self.settings.INGESTION_LEASE_SECONDS,
        )
        if not job:
            return False
        if job.get("kind") == DURABLE_REINDEX_KIND:
            await self._process_claimed_reindex_job(job)
        else:
            await self._process_claimed_durable_job(job)
        return True

    async def _process_claimed_reindex_job(self, job: dict) -> None:
        job_id = str(job.get("job_id") or "")
        lease_token = str(job.get("lease_token") or "")
        try:
            snapshot = self._durable_snapshot(job)
            document_id = job.get("reindex_document_id")
            if not isinstance(document_id, str) or not document_id:
                raise InvalidIngestionInputError("Re-index target is invalid")
            if not await self.job_store.begin_publish(job_id, lease_token):
                return
            refreshed = await self._refresh_document_embeddings(
                job_id=job_id,
                lease_token=lease_token,
                document_id=document_id,
                snapshot=snapshot,
            )
            if not refreshed:
                raise InvalidIngestionInputError("Knowledge document is no longer available for re-indexing")
            if await self.job_store.complete(job_id, lease_token):
                logger.info("durable_knowledge_reindex_completed", job_id=job_id, chunks=refreshed)
        except Exception as exc:
            permanent = isinstance(exc, InvalidIngestionInputError)
            try:
                await self.job_store.retry_or_fail(
                    job_id,
                    lease_token,
                    error=_public_job_error(exc),
                    retryable=not permanent,
                    max_attempts=max(1, int(self.settings.INGESTION_MAX_ATTEMPTS)),
                    retry_delay_seconds=max(1, int(self.settings.INGESTION_RETRY_DELAY_SECONDS)),
                )
            except JobStoreUnavailableError:
                pass
            logger.warning(
                "durable_knowledge_reindex_failed",
                job_id=job_id,
                error_type=type(exc).__name__,
                retryable=not permanent,
            )

    async def _refresh_document_embeddings(
        self,
        *,
        job_id: str,
        lease_token: str,
        document_id: str,
        snapshot: dict,
    ) -> int:
        """Refresh existing document chunks in place, preserving logical IDs."""
        try:
            collection = connection_manager.get_brand_db(snapshot["brand_slug"])["knowledge_base"]
        except Exception as exc:
            raise IngestionStorageError("Knowledge storage is unavailable") from exc

        refreshed = 0
        found = False
        for query in ({"doc_id": document_id}, {"metadata.job_id": document_id}):
            cursor = collection.find(query)
            async for document in cursor:
                found = True
                if not await self.job_store.renew_lease(
                    job_id, lease_token, lease_seconds=self.settings.INGESTION_LEASE_SECONDS
                ):
                    return refreshed
                content = document.get("content")
                chunk_id = document.get("chunk_id") or document.get("_id")
                if not isinstance(content, str) or not content.strip() or not isinstance(chunk_id, str):
                    raise InvalidIngestionInputError("Knowledge document has an invalid chunk")
                embeddings = await self._generate_embeddings(content)
                self._validate_embedding(embeddings)
                updated_at = _utc_timestamp()
                try:
                    await collection.update_one(
                        {"_id": document.get("_id")},
                        {"$set": {"embeddings": embeddings, "reindexed_at": updated_at, "reindex_job_id": job_id, "updated_at": updated_at}},
                    )
                    if self.qdrant:
                        refreshed_document = dict(document)
                        refreshed_document.update({"embeddings": embeddings, "updated_at": updated_at, "reindexed_at": updated_at, "reindex_job_id": job_id})
                        await self.qdrant.upsert_chunk(refreshed_document, snapshot["brand_slug"])
                except Exception as exc:
                    raise IngestionStorageError("Failed to refresh knowledge document vectors") from exc
                refreshed += 1
                if not await self.job_store.mark_progress(job_id, lease_token, refreshed):
                    return refreshed
            if found:
                break
        return refreshed

    async def _process_claimed_durable_job(self, job: dict) -> None:
        job_id = str(job.get("job_id") or "")
        lease_token = str(job.get("lease_token") or "")
        try:
            snapshot = self._durable_snapshot(job)
            if job.get("phase") != "publishing":
                await self._stage_claimed_job(job, lease_token, snapshot)
                if not await self.job_store.begin_publish(job_id, lease_token):
                    # Cancellation that wins before the publish fence leaves no
                    # visible chunks. Staging/payload retention is then removed.
                    current = await self.job_store.get(job_id)
                    if current and current.get("status") == "cancelled":
                        await self._cleanup_durable_artifacts(job_id)
                    return
            elif not await self.job_store.renew_lease(
                job_id,
                lease_token,
                lease_seconds=self.settings.INGESTION_LEASE_SECONDS,
            ):
                return
            await self._publish_claimed_job(job_id, lease_token, snapshot)
            if await self.job_store.complete(job_id, lease_token):
                await self._cleanup_durable_artifacts(job_id)
                logger.info("durable_ingestion_job_completed", job_id=job_id)
        except Exception as exc:
            permanent = isinstance(exc, (InvalidIngestionInputError, InvalidIngestionPayloadError))
            try:
                await self.job_store.retry_or_fail(
                    job_id,
                    lease_token,
                    error=_public_job_error(exc),
                    retryable=not permanent,
                    max_attempts=max(1, int(self.settings.INGESTION_MAX_ATTEMPTS)),
                    retry_delay_seconds=max(1, int(self.settings.INGESTION_RETRY_DELAY_SECONDS)),
                )
                current = await self.job_store.get(job_id)
                if current and current.get("status") in {"error", "cancelled"}:
                    await self._cleanup_durable_artifacts(job_id)
            except JobStoreUnavailableError:
                # The original failure must remain recoverable from its lease;
                # never turn a Mongo outage into an unsafe local terminal state.
                pass
            logger.warning(
                "durable_ingestion_job_failed",
                job_id=job_id,
                error_type=type(exc).__name__,
                retryable=not permanent,
            )

    def _durable_snapshot(self, job: dict) -> dict:
        snapshot = job.get("snapshot")
        if not isinstance(snapshot, dict):
            raise InvalidIngestionInputError("Ingestion snapshot is missing")
        required = ("brand_id", "brand_slug", "chunk_size", "chunk_overlap")
        if any(snapshot.get(key) in (None, "") for key in required):
            raise InvalidIngestionInputError("Ingestion snapshot is invalid")
        # Scope fields are duplicated outside the snapshot for authorization.
        # Requiring equality prevents a malformed Mongo record from silently
        # writing to a different brand database.
        for key in ("brand_id", "brand_slug"):
            if job.get(key) != snapshot.get(key):
                raise InvalidIngestionInputError("Ingestion scope snapshot is invalid")
        if job.get("agent_id") != snapshot.get("agent_id"):
            raise InvalidIngestionInputError("Ingestion scope snapshot is invalid")
        try:
            chunk_size, chunk_overlap = int(snapshot["chunk_size"]), int(snapshot["chunk_overlap"])
        except (TypeError, ValueError) as exc:
            raise InvalidIngestionInputError("Ingestion chunking snapshot is invalid") from exc
        if chunk_size <= 0 or chunk_overlap < 0 or chunk_overlap >= chunk_size:
            raise InvalidIngestionInputError("Ingestion chunking snapshot is invalid")
        return {
            "agent_id": str(snapshot["agent_id"]) if snapshot.get("agent_id") else None,
            "brand_id": str(snapshot["brand_id"]),
            "brand_slug": str(snapshot["brand_slug"]),
            "chunk_size": chunk_size,
            "chunk_overlap": chunk_overlap,
        }

    async def _stage_claimed_job(self, job: dict, lease_token: str, snapshot: dict) -> None:
        job_id = str(job["job_id"])
        payload_refs = job.get("payload_refs")
        if not isinstance(payload_refs, list) or len(payload_refs) != int(job.get("files_count") or 0):
            raise InvalidIngestionInputError("Ingestion payload references are invalid")
        for index, ref in enumerate(sorted(payload_refs, key=lambda item: item.get("file_index", -1))):
            if not isinstance(ref, dict) or ref.get("file_index") != index or not isinstance(ref.get("payload_id"), str):
                raise InvalidIngestionInputError("Ingestion payload references are invalid")
            if not await self.job_store.renew_lease(
                job_id, lease_token, lease_seconds=self.settings.INGESTION_LEASE_SECONDS
            ):
                return
            file_data = await self.payload_store.load(ref["payload_id"], job_id=job_id)
            try:
                chunks = await self._extract_and_chunk(
                    file_data["content"],
                    file_data["content_type"],
                    file_data["filename"],
                    chunk_size=snapshot["chunk_size"],
                    chunk_overlap=snapshot["chunk_overlap"],
                )
            except (UnicodeDecodeError, ValueError, json.JSONDecodeError) as exc:
                raise InvalidIngestionInputError("Document cannot be processed") from exc
            chunks = self._apply_payload_context(
                chunks,
                context=file_data.get("context"),
                filename=file_data["filename"],
                content_type_header=file_data["content_type"],
            )
            if not chunks:
                raise InvalidIngestionInputError("Document contains no ingestible content")
            for chunk_index, chunk in enumerate(chunks):
                if not await self.job_store.renew_lease(
                    job_id, lease_token, lease_seconds=self.settings.INGESTION_LEASE_SECONDS
                ):
                    return
                await self._stage_chunk(
                    job_id=job_id,
                    file_index=index,
                    chunk_index=chunk_index,
                    chunk=chunk,
                    filename=file_data["filename"],
                    snapshot=snapshot,
                )
            if not await self.job_store.mark_progress(job_id, lease_token, index + 1):
                return

    @staticmethod
    def _normalize_folder_path(value: object) -> str:
        raw_path = str(value or "/").strip()
        if not raw_path:
            raw_path = "/"
        if not raw_path.startswith("/"):
            raw_path = f"/{raw_path}"
        normalized = re.sub(r"/+", "/", raw_path).rstrip("/")
        return normalized or "/"

    @staticmethod
    def _source_type(content_type: str, filename: str) -> str:
        normalized_content_type = (content_type or "").split(";", 1)[0].strip().lower()
        content_types = {
            "application/pdf": "pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
            "text/plain": "txt",
            "text/markdown": "md",
            "text/html": "html",
            "application/json": "json",
            "text/csv": "csv",
        }
        if normalized_content_type in content_types:
            return content_types[normalized_content_type]
        lowered_filename = (filename or "").lower()
        for suffix, source_type in {
            ".pdf": "pdf", ".docx": "docx", ".txt": "txt", ".md": "md",
            ".markdown": "md", ".html": "html", ".htm": "html", ".json": "json", ".csv": "csv",
        }.items():
            if lowered_filename.endswith(suffix):
                return source_type
        return "unknown"

    def _apply_payload_context(
        self,
        chunks: list[dict],
        *,
        context: object,
        filename: str,
        content_type_header: str,
    ) -> list[dict]:
        """Apply encrypted knowledge metadata after source extraction.

        The job document contains only opaque references. This method is the
        first point at which optional folder/product/dealer metadata is made
        available to the worker, after its encrypted payload is recovered.
        """
        if context is None:
            return chunks
        if not isinstance(context, dict):
            raise InvalidIngestionInputError("Document upload context is invalid")
        allowed = {"kb_content_type", "folder_path", "product_data", "dealer_data"}
        if set(context) - allowed:
            raise InvalidIngestionInputError("Document upload context is invalid")

        kb_content_type = context.get("kb_content_type")
        if kb_content_type is not None and (
            not isinstance(kb_content_type, str)
            or kb_content_type not in {"product", "dealer", "faq", "office", "category", "guide", "document"}
        ):
            raise InvalidIngestionInputError("Document content type is invalid")
        product_data = context.get("product_data")
        dealer_data = context.get("dealer_data")
        if product_data is not None and not isinstance(product_data, dict):
            raise InvalidIngestionInputError("Product metadata is invalid")
        if dealer_data is not None and not isinstance(dealer_data, dict):
            raise InvalidIngestionInputError("Dealer metadata is invalid")
        if product_data is not None and dealer_data is not None:
            raise InvalidIngestionInputError("Document cannot be both product and dealer data")

        folder = self._normalize_folder_path(context.get("folder_path"))
        name = filename.strip().strip("/") or "untitled"
        path = f"{folder.rstrip('/')}/{name}" if folder != "/" else f"/{name}"
        source_type = self._source_type(content_type_header, filename)
        for chunk in chunks:
            metadata = dict(chunk.get("metadata") or {})
            metadata.update(
                {
                    "filename": filename,
                    "name": name,
                    "folder": folder,
                    "path": path,
                    "content_type_header": content_type_header,
                    "source_type": source_type,
                }
            )
            chunk["metadata"] = metadata
            if kb_content_type is not None:
                chunk["content_type"] = kb_content_type
            if product_data is not None:
                chunk["product_data"] = dict(product_data)
                chunk["dealer_data"] = None
            if dealer_data is not None:
                chunk["dealer_data"] = dict(dealer_data)
                chunk["product_data"] = None
        return chunks

    @staticmethod
    def _durable_chunk_id(job_id: str, file_index: int, chunk_index: int) -> str:
        digest = hashlib.sha256(f"{job_id}:{file_index}:{chunk_index}".encode("utf-8")).hexdigest()
        return f"ingest-{digest}"

    def _staged_collection(self):
        try:
            return connection_manager.get_system_db()[STAGED_CHUNKS_COLLECTION]
        except (RuntimeError, AttributeError, KeyError, TypeError) as exc:
            raise IngestionStorageError("Staging storage is unavailable") from exc

    async def _ensure_staged_indexes(self, collection: Any) -> None:
        # Index creation is idempotent at Mongo; keeping it local makes workers
        # deployable without requiring a one-time migration command.
        try:
            await collection.create_index("expires_at", expireAfterSeconds=0, name="ingestion_staged_chunks_expiry_idx")
            await collection.create_index([("job_id", 1), ("file_index", 1), ("chunk_index", 1)], name="ingestion_staged_chunks_job_idx")
        except Exception as exc:
            logger.warning("ingestion_staging_index_setup_failed", error_type=type(exc).__name__)

    async def _stage_chunk(
        self,
        *,
        job_id: str,
        file_index: int,
        chunk_index: int,
        chunk: dict,
        filename: str,
        snapshot: dict,
    ) -> None:
        content = chunk.get("content")
        if not isinstance(content, str) or not content.strip():
            raise InvalidIngestionInputError("Document contains invalid text")
        embeddings = await self._generate_embeddings(content)
        chunk_id = self._durable_chunk_id(job_id, file_index, chunk_index)
        metadata = dict(chunk.get("metadata") or {})
        metadata.update(
            {
                "job_id": job_id,
                "agent_id": snapshot["agent_id"],
                "brand_id": snapshot["brand_id"],
                "brand_slug": snapshot["brand_slug"],
                "ingestion_file_index": file_index,
                "ingestion_chunk_index": chunk_index,
            }
        )
        now = _utc_timestamp()
        document = {
            "_id": chunk_id,
            "chunk_id": chunk_id,
            "job_id": job_id,
            "file_index": file_index,
            "chunk_index": chunk_index,
            "filename": filename,
            "content": content,
            "embeddings": embeddings,
            "metadata": metadata,
            "content_type": chunk.get("content_type") or "guide",
            "product_data": chunk.get("product_data"),
            "dealer_data": chunk.get("dealer_data"),
            "agent_id": snapshot["agent_id"],
            "brand_id": snapshot["brand_id"],
            "brand_slug": snapshot["brand_slug"],
            "updated_at": now,
            "expires_at": datetime.now(timezone.utc) + timedelta(seconds=self.settings.INGESTION_JOB_TTL_SECONDS),
        }
        collection = self._staged_collection()
        await self._ensure_staged_indexes(collection)
        try:
            update_doc = dict(document)
            update_doc.pop("_id")
            await collection.update_one(
                {"_id": chunk_id},
                {"$set": update_doc, "$setOnInsert": {"created_at": now}},
                upsert=True,
            )
        except Exception as exc:
            logger.warning("ingestion_chunk_stage_failed", job_id=job_id, error_type=type(exc).__name__)
            raise IngestionStorageError("Unable to stage document chunk") from exc

    async def _publish_claimed_job(self, job_id: str, lease_token: str, snapshot: dict) -> None:
        published_count = 0
        async for staged in self._iter_staged_chunks(job_id):
            if not await self.job_store.renew_lease(
                job_id, lease_token, lease_seconds=self.settings.INGESTION_LEASE_SECONDS
            ):
                return
            await self._publish_staged_chunk(staged, snapshot)
            published_count += 1
        if not published_count:
            raise IngestionStorageError("No staged chunks are available for publishing")

    async def _iter_staged_chunks(self, job_id: str) -> AsyncIterator[dict]:
        collection = self._staged_collection()
        try:
            cursor = collection.find({"job_id": job_id})
            if hasattr(cursor, "sort"):
                cursor = cursor.sort([("file_index", 1), ("chunk_index", 1)])
            async for document in cursor:
                yield document
        except IngestionStorageError:
            raise
        except Exception as exc:
            logger.warning("ingestion_staged_chunk_read_failed", job_id=job_id, error_type=type(exc).__name__)
            raise IngestionStorageError("Unable to read staged document chunks") from exc

    async def _publish_staged_chunk(self, staged: dict, snapshot: dict) -> None:
        chunk_id = staged.get("chunk_id")
        if not isinstance(chunk_id, str) or not chunk_id:
            raise IngestionStorageError("Staged chunk is invalid")
        self._validate_embedding(staged.get("embeddings"))
        # Deliberately use the creation-time slug. Do not resolve current agent
        # ownership here: an agent move must not change a queued job's target.
        try:
            brand_db = connection_manager.get_brand_db(snapshot["brand_slug"])
            collection = brand_db["knowledge_base"]
            document = {
                "_id": chunk_id,
                "chunk_id": chunk_id,
                "doc_id": staged.get("job_id"),
                "job_id": staged.get("job_id"),
                "filename": staged.get("filename"),
                "content": staged.get("content"),
                "embeddings": staged.get("embeddings"),
                "metadata": dict(staged.get("metadata") or {}),
                "content_type": staged.get("content_type") or "guide",
                "product_data": staged.get("product_data"),
                "dealer_data": staged.get("dealer_data"),
                "agent_id": snapshot["agent_id"],
                "brand_id": snapshot["brand_id"],
                "brand_slug": snapshot["brand_slug"],
                "updated_at": _utc_timestamp(),
            }
            update_doc = dict(document)
            update_doc.pop("_id")
            await collection.update_one(
                {"_id": chunk_id},
                {"$set": update_doc, "$setOnInsert": {"created_at": _utc_timestamp()}},
                upsert=True,
            )
            if self.qdrant:
                await self.qdrant.upsert_chunk(document, snapshot["brand_slug"])
        except IngestionStorageError:
            raise
        except Exception as exc:
            logger.warning("ingestion_chunk_publish_failed", chunk_id=chunk_id, error_type=type(exc).__name__)
            raise IngestionStorageError("Failed to store knowledge-base document") from exc

    async def _cleanup_durable_artifacts(self, job_id: str) -> None:
        await self.payload_store.delete_job_payloads(job_id)
        try:
            await self._staged_collection().delete_many({"job_id": job_id})
        except Exception as exc:
            logger.warning("ingestion_staging_cleanup_failed", job_id=job_id, error_type=type(exc).__name__)

    async def process_documents(self, job_id: str, files: List[dict], agent_id: Optional[str] = None):
        """Process documents in background."""
        try:
            await self.job_store.update(job_id, {"status": "processing"})

            chunk_size, chunk_overlap = await self._resolve_chunking(agent_id)

            for i, file_data in enumerate(files):
                # Extract file information
                content = file_data['content']
                filename = file_data['filename']
                content_type = file_data['content_type']

                # Process based on content type
                chunks = await self._extract_and_chunk(
                    content, content_type, filename,
                    chunk_size=chunk_size, chunk_overlap=chunk_overlap,
                )

                # Store chunks with embeddings
                await self._store_chunks(chunks, job_id, filename, agent_id)

                # Update progress
                await self.job_store.update(job_id, {"processed_count": i + 1})

                logger.info("Processed file", job_id=job_id, filename=filename, chunks_count=len(chunks))

            await self.job_store.update(
                job_id,
                {"status": "completed", "completed_at": _utc_timestamp()},
            )
            logger.info("Completed ingestion job", job_id=job_id)

        except Exception as exc:
            await self.job_store.update(
                job_id,
                {
                    "status": "error",
                    "error": _public_job_error(exc),
                    "completed_at": _utc_timestamp(),
                },
            )
            logger.error(
                "ingestion_job_processing_failed",
                job_id=job_id,
                error_type=type(exc).__name__,
            )
    
    async def process_chunk(
        self,
        request: IngestionRequest,
        *,
        agent_id: Optional[str] = None,
    ) -> IngestionResponse:
        """Process a single text chunk."""
        try:
            # Generate embeddings
            embeddings = await self._generate_embeddings(request.text)
            
            # Create chunk document
            metadata = dict(request.metadata or {})
            if agent_id:
                # The route has already authorized this agent. Persist that
                # canonical scope rather than trusting caller-supplied metadata.
                metadata["agent_id"] = agent_id

            chunk_doc = {
                "doc_id": request.doc_id or str(uuid.uuid4()),
                "content": request.text,
                "embeddings": embeddings,
                "metadata": metadata,
                "agent_id": agent_id,
                "created_at": datetime.utcnow().isoformat()
            }
            
            # Store in vector database
            chunk_id = await self._store_chunk(chunk_doc)
            
            return IngestionResponse(
                success=True,
                chunk_id=chunk_id,
                message="Chunk processed successfully"
            )
            
        except Exception as exc:
            logger.error("ingestion_chunk_processing_failed", error_type=type(exc).__name__)
            return IngestionResponse(
                success=False,
                chunk_id=None,
                message="Unable to process chunk",
            )
    
    async def get_job_status(self, job_id: str) -> Optional[IngestionStatus]:
        """Get the status of an ingestion job."""
        job_info = await self.job_store.get(job_id)
        if not job_info:
            return None

        internal_status = job_info["status"]
        public_status = {
            "queued": "pending",
            "running": "processing",
            "publishing": "processing",
        }.get(internal_status, internal_status)
        return IngestionStatus(
            job_id=job_id,
            status=public_status,
            files_count=job_info["files_count"],
            processed_count=job_info["processed_count"],
            error=_safe_stored_job_error(public_status, job_info.get("error")),
            created_at=job_info.get("created_at"),
            completed_at=job_info.get("completed_at"),
        )

    async def cancel_job(self, job_id: str) -> bool:
        """Cancel an ingestion job."""
        job_info = await self.job_store.get(job_id)
        if not job_info:
            return False
        if hasattr(self.job_store, "cancel"):
            cancelled = await self.job_store.cancel(job_id)
        else:
            cancelled = await self.job_store.update(
                job_id,
                {"status": "cancelled", "completed_at": _utc_timestamp()},
            )
        if cancelled:
            logger.info("ingestion_job_cancelled", job_id=job_id)
        return bool(cancelled)
    
    async def get_documents(self, agent_id: Optional[str] = None) -> List[dict]:
        """Get uploaded documents, optionally filtered by agent_id."""
        try:
            # Get brand-specific database
            if agent_id:
                brand_db = await connection_manager.get_brand_db_by_agent_id(agent_id)
            else:
                # Fallback to system database (though documents should be in brand databases)
                brand_db = connection_manager.get_system_db()
            
            chunks_collection = brand_db["knowledge_base"]
            
            # Build query (no need to filter by agent_id since we're in brand-specific DB)
            query = {}
            if agent_id and not hasattr(connection_manager, 'get_brand_db_by_agent_id'):
                # Fallback for legacy compatibility
                query["agent_id"] = agent_id
            
            # Aggregate to group by filename and get metadata
            pipeline = [
                {"$match": query},
                {
                    "$group": {
                        "_id": "$filename",
                        "agent_id": {"$first": "$agent_id"},
                        "job_id": {"$first": "$job_id"},
                        "chunks_count": {"$sum": 1},
                        "created_at": {"$first": "$created_at"},
                        "content_type": {"$first": "$metadata.content_type"}
                    }
                },
                {"$sort": {"created_at": -1}}
            ]
            
            cursor = chunks_collection.aggregate(pipeline)
            documents = []
            
            async for doc in cursor:
                documents.append({
                    "filename": doc["_id"],
                    "agent_id": doc.get("agent_id"),
                    "job_id": doc.get("job_id"),
                    "chunks_count": doc["chunks_count"],
                    "created_at": doc.get("created_at"),
                    "content_type": doc.get("content_type")
                })
            
            logger.info("Retrieved documents", count=len(documents), agent_id=agent_id)
            return documents
            
        except Exception as e:
            logger.error("Error getting documents from MongoDB", error=str(e))
            return []

    
    def _extract_product_data(self, json_obj: dict) -> Optional[dict]:
        """Extract product fields from JSON object."""
        try:
            # Check if this looks like product data
            has_product_fields = any(
                field in json_obj 
                for field in ['sku', 'name', 'price', 'product_id', 'product_name']
            )
            
            if not has_product_fields:
                return None
            
            product = {
                "sku": json_obj.get("sku") or json_obj.get("product_id"),
                "name": json_obj.get("name") or json_obj.get("product_name"),
                "price": json_obj.get("price"),
                "currency": json_obj.get("currency"),
                "category": json_obj.get("category"),
                "image_url": json_obj.get("image_url") or json_obj.get("image"),
                "product_url": json_obj.get("product_url") or json_obj.get("url"),
                "in_stock": json_obj.get("in_stock", True),
                "features": json_obj.get("features", []),
            }
            # Keep commerce variant and provenance fields intact. Product cards
            # and Shopify reconciliation depend on these, so durable bulk
            # ingestion must not silently flatten them into generic text.
            for field in (
                "id", "price_unit", "currency_source", "product_group_id", "handle", "parent_name",
                "has_variants", "variant_count", "price_min", "price_max", "default_variant_id",
                "variant_id", "variant_sku", "variant_title", "variant_options", "variant_url", "variants",
                "source_type", "source_url", "source_product_id", "source_variant_id", "source_key", "source_active",
            ):
                if field in json_obj:
                    product[field] = json_obj.get(field)
            return product
        except Exception as e:
            logger.warning("Failed to extract product data", error=str(e))
            return None
    
    def _extract_dealer_data(self, json_obj: dict) -> Optional[dict]:
        """Extract dealer fields from JSON object."""
        try:
            # Check if this looks like dealer data
            has_dealer_fields = any(
                field in json_obj 
                for field in ['dealer_id', 'name', 'city', 'phone']
            )
            
            if not has_dealer_fields:
                return None
            
            return {
                "dealer_id": json_obj.get("dealer_id"),
                "name": json_obj.get("name"),
                "city": json_obj.get("city"),
                "state": json_obj.get("state"),
                "phone": json_obj.get("phone"),
                "email": json_obj.get("email"),
                "address": json_obj.get("address")
            }
        except Exception as e:
            logger.warning("Failed to extract dealer data", error=str(e))
            return None
    
    def _detect_content_type(self, json_obj: dict) -> str:
        """Auto-detect content type from JSON structure."""
        # Check for product fields
        if any(f in json_obj for f in ['sku', 'product_id', 'price']):
            return "product"
        
        # Check for dealer fields
        if any(f in json_obj for f in ['dealer_id', 'city', 'phone']):
            return "dealer"
        
        # Check for FAQ fields
        if 'question' in json_obj or 'answer' in json_obj:
            return "faq"
        
        # Check for office fields
        if 'office_id' in json_obj or 'office_name' in json_obj:
            return "office"
        
        # Check for category fields
        if 'category_id' in json_obj or 'category_name' in json_obj:
            return "category"
        
        # Default to guide
        return "guide"

    async def _resolve_chunking(self, agent_id: Optional[str]) -> Tuple[int, int]:
        """Resolve chunk size/overlap via the shared chunking module."""
        return await resolve_agent_chunking(agent_id)

    def snapshot_chunking_from_agent(self, agent: dict) -> Tuple[int, int]:
        """Freeze chunking from the already-authorized creation-time agent row."""
        chunking = (((agent.get("configuration") or {}).get("rag") or {}).get("chunking") or {})
        try:
            size = int(chunking.get("chunk_size", DEFAULT_CHUNK_SIZE))
            overlap = int(chunking.get("chunk_overlap", DEFAULT_CHUNK_OVERLAP))
        except (TypeError, ValueError):
            size, overlap = DEFAULT_CHUNK_SIZE, DEFAULT_CHUNK_OVERLAP
        return clamp_chunking(size, overlap)

    def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from a PDF using pypdf."""
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError("PDF extraction requires pypdf to be installed") from exc

        reader = PdfReader(io.BytesIO(content))
        page_text = []
        for page_number, page in enumerate(reader.pages, start=1):
            extracted = (page.extract_text() or "").strip()
            if extracted:
                page_text.append(f"Page {page_number}\n{extracted}")
        return "\n\n".join(page_text)

    def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from a DOCX using python-docx."""
        self._validate_docx_archive(content)
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("DOCX extraction requires python-docx to be installed") from exc

        document = Document(io.BytesIO(content))
        paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        table_rows = []
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_rows.append(" | ".join(cells))
        return "\n\n".join(paragraphs + table_rows)

    def _validate_docx_archive(self, content: bytes) -> None:
        """Reject compressed DOCX archives that can exhaust worker memory/CPU."""
        max_files = max(1, int(self.settings.MAX_ARCHIVE_FILES))
        max_uncompressed = max(1, int(self.settings.MAX_ARCHIVE_UNCOMPRESSED_SIZE_MB)) * 1024 * 1024
        max_ratio = max(1, int(self.settings.MAX_ARCHIVE_COMPRESSION_RATIO))
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as archive:
                entries = archive.infolist()
                if len(entries) > max_files:
                    raise InvalidIngestionInputError("DOCX archive contains too many files")
                uncompressed = sum(max(0, int(entry.file_size)) for entry in entries)
                compressed = sum(max(0, int(entry.compress_size)) for entry in entries)
        except InvalidIngestionInputError:
            raise
        except (zipfile.BadZipFile, OSError, ValueError) as exc:
            raise InvalidIngestionInputError("DOCX archive is invalid") from exc
        if uncompressed > max_uncompressed:
            raise InvalidIngestionInputError("DOCX archive exceeds the uncompressed size limit")
        # Empty archives are invalid DOCX documents and should be rejected by
        # python-docx. Avoid division by zero while still catching a compressed
        # payload that expands beyond its declared ratio.
        if uncompressed and (compressed == 0 or uncompressed / max(compressed, 1) > max_ratio):
            raise InvalidIngestionInputError("DOCX archive compression ratio is too high")

    async def _extract_and_chunk(
        self,
        content: bytes,
        content_type: str,
        filename: str,
        chunk_size: int = DEFAULT_CHUNK_SIZE,
        chunk_overlap: int = DEFAULT_CHUNK_OVERLAP,
    ) -> List[dict]:
        """Extract text and create chunks with structured data extraction."""
        # Extract text based on content type
        json_data = None
        structured_content_type = None
        lower_name = (filename or "").lower()

        if content_type == "application/json":
            text = content.decode('utf-8')
            # For JSON files, try to extract structured data
            try:
                json_data = json.loads(text)
                
                # If it's a list, process each item
                if isinstance(json_data, list):
                    # Process the first item to detect type
                    if json_data and isinstance(json_data[0], dict):
                        structured_content_type = self._detect_content_type(json_data[0])
                        logger.info(
                            "Detected content type from JSON",
                            filename=filename,
                            detected_type=structured_content_type,
                            items_count=len(json_data)
                        )
                
                # Convert JSON to readable text format
                text = self._json_to_text(json_data)
            except json.JSONDecodeError:
                pass  # Treat as plain text
        elif content_type in ("text/plain", "text/markdown", "text/csv", "text/html"):
            text = content.decode('utf-8', errors='replace')
        elif content_type == "application/pdf" or lower_name.endswith(".pdf"):
            text = self._extract_pdf_text(content)
        elif (
            content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            or lower_name.endswith(".docx")
        ):
            text = self._extract_docx_text(content)
        else:
            # For other types, try to decode as text but reject binary garbage
            # instead of embedding unreadable bytes.
            text = content.decode('utf-8', errors='replace')
            if text and (text.count('�') / len(text)) > 0.2:
                raise ValueError(
                    f"File {filename} ({content_type or 'unknown type'}) does not contain extractable text"
                )

        # Chunk the text (bounds are enforced by the shared chunker)
        chunks = []
        
        # If we have structured JSON data (list of items), create one chunk per item
        if json_data and isinstance(json_data, list) and structured_content_type:
            for idx, item in enumerate(json_data):
                if not isinstance(item, dict):
                    continue
                
                # Extract structured data based on content type
                product_data = None
                dealer_data = None
                
                if structured_content_type == "product":
                    product_data = self._extract_product_data(item)
                elif structured_content_type == "dealer":
                    dealer_data = self._extract_dealer_data(item)
                
                # Create readable text for this item
                item_text = self._json_to_text(item)
                
                chunks.append({
                    "content": item_text,
                    "content_type": structured_content_type,
                    "product_data": product_data,
                    "dealer_data": dealer_data,
                    "metadata": {
                        "filename": filename,
                        "chunk_index": idx,
                        "content_type": content_type,
                        "item_index": idx,
                        "structured_type": structured_content_type
                    }
                })
        else:
            # Standard chunking for non-structured content via the shared
            # paragraph-aware chunker (same chunks as the Knowledge Base upload path).
            for chunk_index, piece in enumerate(chunk_text(text, chunk_size=chunk_size, chunk_overlap=chunk_overlap)):
                chunks.append({
                    "content": piece,
                    "content_type": "guide",  # Default content type
                    "product_data": None,
                    "dealer_data": None,
                    "metadata": {
                        "filename": filename,
                        "chunk_index": chunk_index,
                        "content_type": content_type,
                    }
                })

        return chunks
    
    def _json_to_text(self, data, prefix="") -> str:
        """Convert JSON data to readable text."""
        lines = []
        
        if isinstance(data, dict):
            for key, value in data.items():
                if isinstance(value, (dict, list)):
                    lines.append(f"{prefix}{key}:")
                    lines.append(self._json_to_text(value, prefix + "  "))
                else:
                    lines.append(f"{prefix}{key}: {value}")
        elif isinstance(data, list):
            for i, item in enumerate(data):
                if isinstance(item, (dict, list)):
                    lines.append(f"{prefix}[{i}]:")
                    lines.append(self._json_to_text(item, prefix + "  "))
                else:
                    lines.append(f"{prefix}- {item}")
        else:
            lines.append(f"{prefix}{data}")
        
        return "\n".join(lines)
    
    async def _generate_embeddings(self, text: str) -> List[float]:
        """Generate embeddings using Voyage AI."""
        voyage_config = await self._get_voyage_runtime_config()
        api_key = voyage_config["api_key"]
        base_url = voyage_config["base_url"].rstrip("/")
        model = voyage_config["model"]

        if not api_key:
            logger.error("voyage_api_key_not_configured", text_length=len(text))
            raise IngestionEmbeddingError("Voyage embedding API key is not configured")

        try:
            async with httpx.AsyncClient() as client:
                response = await client.post(
                    f"{base_url}/embeddings",
                    headers={
                        "Authorization": f"Bearer {api_key}",
                        "Content-Type": "application/json"
                    },
                    json={
                        "input": [text],
                        "model": model
                    },
                    timeout=30.0
                )
        except httpx.HTTPError as exc:
            logger.error("voyage_embedding_request_failed", error_type=type(exc).__name__)
            raise IngestionEmbeddingError("Voyage embedding request failed") from exc

        if response.status_code != 200:
            logger.error("voyage_embedding_request_rejected", status_code=response.status_code)
            raise IngestionEmbeddingError(
                f"Voyage embedding request failed with HTTP {response.status_code}"
            )

        try:
            embeddings = response.json()["data"][0]["embedding"]
        except (IndexError, KeyError, TypeError, ValueError) as exc:
            logger.error("voyage_embedding_response_invalid", error_type=type(exc).__name__)
            raise IngestionEmbeddingError("Voyage returned an invalid embeddings response") from exc

        self._validate_embedding(embeddings)
        logger.debug("Generated embeddings", dimensions=len(embeddings))
        return embeddings

    def _validate_embedding(self, embeddings: object) -> None:
        """Reject malformed or placeholder vectors before any storage write."""
        if not isinstance(embeddings, list):
            raise IngestionEmbeddingError("Voyage returned an embedding with an invalid type")

        if len(embeddings) != self.settings.VECTOR_DIMENSIONS:
            raise IngestionEmbeddingError(
                "Voyage returned an embedding with unexpected dimensions "
                f"({len(embeddings)}; expected {self.settings.VECTOR_DIMENSIONS})"
            )

        if any(
            not isinstance(value, (int, float))
            or isinstance(value, bool)
            or not math.isfinite(value)
            for value in embeddings
        ):
            raise IngestionEmbeddingError("Voyage returned an embedding with invalid values")

        if not any(embeddings):
            raise IngestionEmbeddingError("Voyage returned an all-zero embedding")
    
    async def _store_chunks(self, chunks: List[dict], job_id: str, filename: str, agent_id: Optional[str] = None):
        """Store chunks in vector database with structured data."""
        for chunk in chunks:
            # Generate embeddings for this chunk
            chunk["embeddings"] = await self._generate_embeddings(chunk["content"])
            chunk["chunk_id"] = str(uuid.uuid4())  # Generate unique chunk ID
            chunk["job_id"] = job_id
            chunk["filename"] = filename
            chunk["agent_id"] = agent_id
            chunk["created_at"] = datetime.utcnow().isoformat()
            
            # content_type, product_data, dealer_data are already in chunk from _extract_and_chunk()
            # Just ensure they exist
            if "content_type" not in chunk:
                chunk["content_type"] = "guide"
            
            if "product_data" not in chunk:
                chunk["product_data"] = None
            
            if "dealer_data" not in chunk:
                chunk["dealer_data"] = None

            # Ensure metadata object exists and include brand info when possible
            if "metadata" not in chunk or not isinstance(chunk["metadata"], dict):
                chunk.setdefault("metadata", {})

            # If agent_id is provided, try to resolve the agent -> brand and include both
            if agent_id:
                try:
                    system_db = connection_manager.get_system_db()
                    agent_doc = await system_db.agents.find_one({"id": agent_id})
                    if agent_doc:
                        # include both brand_id and brand_slug for cross-referencing
                        if agent_doc.get("brand_id"):
                            chunk["metadata"]["brand_id"] = agent_doc.get("brand_id")
                        if agent_doc.get("brand_slug"):
                            chunk["metadata"]["brand_slug"] = agent_doc.get("brand_slug")
                        # Also ensure agent_id present in metadata for convenience
                        chunk["metadata"]["agent_id"] = agent_id
                except Exception:
                    # If resolving fails, continue without brand info (will fallback)
                    logger.debug("Could not resolve agent -> brand for chunk metadata", agent_id=agent_id)

            # Store in MongoDB
            await self._store_chunk(chunk)
    
    async def _store_chunk(self, chunk_doc: dict) -> str:
        """Store a single chunk in brand-specific MongoDB database."""
        try:
            self._validate_embedding(chunk_doc.get("embeddings"))

            # Get brand-specific database
            agent_id = chunk_doc.get("agent_id")
            if agent_id:
                brand_db = await connection_manager.get_brand_db_by_agent_id(agent_id)
            else:
                # Fallback to system database if no agent_id
                brand_db = connection_manager.get_system_db()
            
            # Get or create the chunks collection in brand database  
            chunks_collection = brand_db["knowledge_base"]
            
            # Add unique ID if not present
            if "_id" not in chunk_doc:
                chunk_doc["_id"] = str(uuid.uuid4())
            
            # Insert the chunk
            result = await chunks_collection.insert_one(chunk_doc)
            inserted_id = getattr(result, "inserted_id", None)
            if inserted_id is None:
                raise IngestionStorageError("MongoDB did not return an inserted chunk ID")

            chunk_id = str(inserted_id)
            if self.qdrant:
                brand_slug = (chunk_doc.get("metadata") or {}).get("brand_slug")
                await self.qdrant.upsert_chunk(chunk_doc, brand_slug)
            
            logger.debug("Stored chunk in MongoDB", chunk_id=chunk_id, agent_id=chunk_doc.get("agent_id"))
            return chunk_id

        except IngestionStorageError:
            raise
        except Exception as exc:
            logger.error("chunk_storage_failed", error_type=type(exc).__name__)
            raise IngestionStorageError("Failed to store knowledge-base chunk") from exc
