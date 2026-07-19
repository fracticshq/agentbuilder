"""Stateless document type detection and text extraction for knowledge uploads."""

from __future__ import annotations

import csv
import html
import io
import json
import re
from typing import Optional

MIME_SOURCE_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt",
    "text/markdown": "md",
    "text/x-markdown": "md",
    "text/html": "html",
    "application/xhtml+xml": "html",
    "application/json": "json",
    "text/json": "json",
    "text/csv": "csv",
    "application/csv": "csv",
    "application/vnd.ms-excel": "csv",
}

EXTENSION_SOURCE_TYPES = {
    ".pdf": "pdf", ".docx": "docx", ".txt": "txt", ".md": "md",
    ".markdown": "md", ".html": "html", ".htm": "html", ".json": "json", ".csv": "csv",
}


class KnowledgeTextExtractionError(RuntimeError):
    """A supported source could not be converted to text."""


def detect_source_type(content_type: Optional[str], filename: Optional[str]) -> Optional[str]:
    normalized_content_type = (content_type or "").split(";", 1)[0].strip().lower()
    if normalized_content_type in MIME_SOURCE_TYPES:
        return MIME_SOURCE_TYPES[normalized_content_type]
    normalized_filename = (filename or "").lower()
    for extension, source_type in EXTENSION_SOURCE_TYPES.items():
        if normalized_filename.endswith(extension):
            return source_type
    return None


def decode_text(content: bytes) -> str:
    return content.decode("utf-8-sig", errors="replace")


def extract_html_text(content: bytes) -> str:
    text = decode_text(content)
    text = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", text)
    text = re.sub(r"(?s)<br\s*/?>", "\n", text)
    text = re.sub(r"(?s)</(p|div|section|article|li|tr|h[1-6])>", "\n", text)
    text = re.sub(r"<[^>]+>", " ", text)
    text = html.unescape(text)
    return re.sub(r"[ \t]+\n", "\n", re.sub(r"\n{3,}", "\n\n", text)).strip()


def extract_csv_text(content: bytes) -> str:
    text = decode_text(content)
    rows = list(csv.reader(io.StringIO(text)))
    if not rows:
        return ""
    headers = [header.strip() for header in rows[0]]
    try:
        has_headers = csv.Sniffer().has_header(text[:2048])
    except csv.Error:
        has_headers = any(headers) and len(rows) > 1
    if len(rows) == 1:
        has_headers = False
    extracted_rows = []
    for index, row in enumerate(rows[1:] if has_headers else rows, start=1):
        if has_headers:
            values = []
            for column_index, value in enumerate(row):
                header = headers[column_index] if column_index < len(headers) and headers[column_index] else f"Column {column_index + 1}"
                values.append(f"{header}: {value.strip()}")
            extracted_rows.append(f"Row {index}\n" + "\n".join(values))
        else:
            extracted_rows.append(f"Row {index}: " + ", ".join(value.strip() for value in row))
    return "\n\n".join(extracted_rows)


def extract_pdf_text(content: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency packaging guard
        raise KnowledgeTextExtractionError("PDF extraction requires pypdf") from exc
    reader = PdfReader(io.BytesIO(content))
    pages = []
    for page_number, page in enumerate(reader.pages, start=1):
        extracted = (page.extract_text() or "").strip()
        if extracted:
            pages.append(f"Page {page_number}\n{extracted}")
    return "\n\n".join(pages)


def extract_docx_text(content: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency packaging guard
        raise KnowledgeTextExtractionError("DOCX extraction requires python-docx") from exc
    document = Document(io.BytesIO(content))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_rows = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_rows.append(" | ".join(cells))
    return "\n\n".join(paragraphs + table_rows)


def extract_text(content: bytes, content_type: str, filename: str) -> str:
    try:
        source_type = detect_source_type(content_type, filename)
        if source_type in {"txt", "md"}:
            return decode_text(content)
        if source_type == "html":
            return extract_html_text(content)
        if source_type == "json":
            return json.dumps(json.loads(decode_text(content)), indent=2)
        if source_type == "csv":
            return extract_csv_text(content)
        if source_type == "pdf":
            return extract_pdf_text(content)
        if source_type == "docx":
            return extract_docx_text(content)
        return decode_text(content)
    except KnowledgeTextExtractionError:
        raise
    except Exception as exc:
        raise KnowledgeTextExtractionError(f"Text extraction failed for {filename}") from exc
