"""
Knowledge Service - Enhanced document processing with structured metadata
Handles both document uploads and bulk JSON imports for products/dealers
"""

import csv
import html
import io
import uuid
import json
import re
from typing import List, Optional, Dict, Any
from datetime import datetime
import structlog
import httpx
from pymongo import UpdateOne

from ..config import Settings
from ..connections import connection_manager
from .chunking import chunk_text, resolve_agent_chunking
from .commerce_config import normalize_commerce_config
from .job_store import JobStore
from .qdrant_vector_service import QdrantVectorService
from .runtime_settings_service import RuntimeSettingsService

logger = structlog.get_logger()


class KnowledgeService:
    """Service for knowledge base operations with structured metadata support."""

    MIME_SOURCE_TYPES = {
        "application/pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "text/plain": "txt",
        "text/markdown": "md",
        "text/x-markdown": "md",
        "text/html": "html",
        "application/xhtml+xml": "html",
        "application/json": "json",
        "text/json": "json",
        "text/csv": "csv",
        "application/csv": "csv",
        "application/vnd.ms-excel": "csv",
    }

    EXTENSION_SOURCE_TYPES = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".txt": "txt",
        ".md": "md",
        ".markdown": "md",
        ".html": "html",
        ".htm": "html",
        ".json": "json",
        ".csv": "csv",
    }
    
    def __init__(self, settings: Settings):
        self.settings = settings
        self.job_store = JobStore()
        self.runtime_settings_service = RuntimeSettingsService(settings)
        self.qdrant = QdrantVectorService(settings) if settings.VECTOR_BACKEND == "qdrant" else None
        self.mongo_client = None
        self.db_cache = {}  # Cache brand-specific databases
        self.brand_scope_cache = {}
        self.db = None
        self.collection = None

    def _normalize_folder_path(self, path: Optional[str]) -> str:
        raw_path = (path or "/").strip()
        if not raw_path:
            raw_path = "/"
        if not raw_path.startswith("/"):
            raw_path = f"/{raw_path}"
        normalized = re.sub(r"/+", "/", raw_path).rstrip("/")
        return normalized or "/"

    def _normalize_item_path(self, folder: Optional[str], name: Optional[str]) -> tuple[str, str, str]:
        clean_folder = self._normalize_folder_path(folder)
        clean_name = (name or "untitled").strip().strip("/") or "untitled"
        return clean_folder, clean_name, f"{clean_folder.rstrip('/')}/{clean_name}" if clean_folder != "/" else f"/{clean_name}"

    def folder_path_from_name(self, name: Optional[str], parent_path: Optional[str] = None) -> str:
        _, clean_name, path = self._normalize_item_path(parent_path, name)
        return path if clean_name else self._normalize_folder_path(parent_path)

    def _normalize_currency(self, value: Any) -> Optional[str]:
        if value is None:
            return None
        currency = str(value).strip()
        return currency.upper() if currency else None

    async def _resolve_configured_default_currency(
        self,
        brand_id: str,
        brand_scope: Optional[Dict[str, Any]] = None,
    ) -> Optional[str]:
        try:
            system_db = connection_manager.get_system_db()
            identifiers = {brand_id}
            if brand_scope:
                for key in ("brand_id", "brand_slug", "identifier"):
                    if brand_scope.get(key):
                        identifiers.add(brand_scope[key])
                identifiers.update(alias for alias in brand_scope.get("aliases", []) if alias)

            agent = await system_db.agents.find_one({
                "$or": [
                    {"id": {"$in": list(identifiers)}},
                    {"brand_id": {"$in": list(identifiers)}},
                    {"brand_slug": {"$in": list(identifiers)}},
                ],
                "configuration.commerce.default_currency": {"$nin": [None, ""]},
            })
            if not agent:
                return None

            commerce = normalize_commerce_config((agent.get("configuration") or {}).get("commerce"))
            return self._normalize_currency(commerce.get("default_currency"))
        except Exception as exc:
            logger.warning("knowledge_default_currency_resolution_failed", brand_id=brand_id, error=str(exc))
            return None

    def _resolve_item_currency(
        self,
        item: Any,
        configured_default_currency: Optional[str] = None,
    ) -> tuple[Optional[str], str]:
        item_currency = self._normalize_currency(getattr(item, "currency", None))
        item_source = getattr(item, "currency_source", None)
        if item_currency:
            return item_currency, item_source if item_source in {"catalog", "configured_default"} else "catalog"

        default_currency = self._normalize_currency(configured_default_currency)
        if default_currency:
            return default_currency, "configured_default"

        return None, "missing"

    def _product_variant_metadata(self, item: Any) -> Dict[str, Any]:
        fields = (
            "product_group_id",
            "handle",
            "parent_name",
            "has_variants",
            "variant_count",
            "price_min",
            "price_max",
            "default_variant_id",
            "variant_id",
            "variant_sku",
            "variant_title",
            "variant_options",
            "variant_url",
        )
        metadata: Dict[str, Any] = {}
        for field in fields:
            value = getattr(item, field, None)
            if value not in (None, "", [], {}):
                metadata[field] = value
        return metadata

    async def _get_knowledge_folders_collection(self, brand_id: str):
        db = await self._get_brand_database(brand_id)
        return db["knowledge_folders"]

    async def _get_voyage_runtime_config(self) -> Dict[str, str]:
        config = await self.runtime_settings_service.get_voyage_runtime_config()
        return {
            "api_key": config["api_key"],
            "base_url": config["base_url"],
            "model": config["model"],
        }
    
    async def _resolve_brand_scope(self, identifier: str) -> Dict[str, Any]:
        """Resolve a brand UUID, brand slug, or agent ID into DB name and query aliases."""
        if identifier in self.brand_scope_cache:
            return self.brand_scope_cache[identifier]

        aliases: List[str] = []

        def add_alias(value: Optional[str]) -> None:
            if value and value not in aliases:
                aliases.append(value)

        add_alias(identifier)
        brand_slug: Optional[str] = None
        canonical_brand_id: Optional[str] = None

        try:
            system_db = connection_manager.get_system_db()

            brand = await system_db.brands.find_one({
                "$or": [
                    {"id": identifier},
                    {"slug": identifier},
                ]
            })
            if brand:
                canonical_brand_id = brand.get("id")
                brand_slug = brand.get("slug") or brand.get("id")
                add_alias(canonical_brand_id)
                add_alias(brand_slug)

            agent = await system_db.agents.find_one({"id": identifier})
            if agent:
                add_alias(agent.get("id"))
                add_alias(agent.get("brand_id"))
                add_alias(agent.get("brand_slug"))
                if not canonical_brand_id:
                    canonical_brand_id = agent.get("brand_id")
                if not brand_slug:
                    brand_slug = agent.get("brand_slug")

                if agent.get("brand_id") or agent.get("brand_slug"):
                    agent_brand = await system_db.brands.find_one({
                        "$or": [
                            {"id": agent.get("brand_id")},
                            {"slug": agent.get("brand_slug")},
                        ]
                    })
                    if agent_brand:
                        canonical_brand_id = agent_brand.get("id") or canonical_brand_id
                        brand_slug = agent_brand.get("slug") or brand_slug
                        add_alias(canonical_brand_id)
                        add_alias(brand_slug)
        except Exception as exc:
            logger.warning("brand_scope_resolution_failed", identifier=identifier, error=str(exc))

        primary_db_key = brand_slug or identifier
        scope = {
            "identifier": identifier,
            "brand_id": canonical_brand_id or identifier,
            "brand_slug": brand_slug,
            "aliases": aliases,
            "db_name": primary_db_key.replace('.', '_')[:63],
        }

        for alias in aliases:
            self.brand_scope_cache[alias] = scope
        self.brand_scope_cache[identifier] = scope
        return scope

    async def _get_brand_database(self, brand_id: str):
        """
        Get brand-specific database.
        
        Accepts a brand UUID, brand slug, or agent ID and resolves to the
        canonical brand database. This keeps old rows visible even when callers
        switch between UUID and slug identifiers.
        """
        # Check if already cached
        if brand_id in self.db_cache:
            return self.db_cache[brand_id]
        
        # Get the MongoDB client from connection manager
        mongo_client = connection_manager.mongodb_client
        if not mongo_client:
            raise RuntimeError("MongoDB not connected")
        
        scope = await self._resolve_brand_scope(brand_id)
        db_name = scope["db_name"]
        self.db_cache[brand_id] = mongo_client[db_name]
        logger.info(
            "brand_database_accessed",
            brand_id=brand_id,
            brand_slug=scope.get("brand_slug"),
            aliases=scope.get("aliases"),
            db_name=db_name,
        )
        
        return self.db_cache[brand_id]

    async def _get_brand_knowledge_collections(self, brand_id: str):
        """Return canonical and legacy knowledge collections for a brand scope."""
        mongo_client = connection_manager.mongodb_client
        if not mongo_client:
            raise RuntimeError("MongoDB not connected")

        scope = await self._resolve_brand_scope(brand_id)
        db_names: List[str] = []

        def add_db_name(value: Optional[str]) -> None:
            if not value:
                return
            db_name = value.replace('.', '_')[:63]
            if db_name not in db_names:
                db_names.append(db_name)

        add_db_name(scope.get("db_name"))
        for alias in scope.get("aliases", []):
            add_db_name(alias)

        return [mongo_client[db_name]["knowledge_base"] for db_name in db_names]
    
    async def _ensure_connection(self, brand_id: Optional[str] = None):
        """Ensure MongoDB connection is established for specific brand."""
        # Use brand-specific database for isolation
        if brand_id:
            self.db = await self._get_brand_database(brand_id)
            self.collection = self.db["knowledge_base"]
        else:
            # Fallback to default database
            self.db = connection_manager.get_mongodb_db()
            self.collection = self.db["knowledge_base"]
    
    # ========================================================================
    # Job Management
    # ========================================================================
    
    async def start_document_upload(
        self,
        content: bytes,
        filename: str,
        content_type_header: str,
        kb_content_type: str,
        brand_id: str,
        agent_id: Optional[str] = None,
        product_data: Optional[Any] = None,
        dealer_data: Optional[Any] = None,
        folder_path: Optional[str] = None,
    ) -> str:
        """Start a document upload job."""
        job_id = str(uuid.uuid4())
        source_type = self.detect_source_type(content_type_header, filename)
        
        await self.job_store.set(job_id, {
            "status": "pending",
            "type": "document",
            "filename": filename,
            "source_type": source_type,
            "content_type_header": content_type_header,
            "content_type": kb_content_type,
            "brand_id": brand_id,
            "agent_id": agent_id,
            "folder": self._normalize_folder_path(folder_path),
            "processed_chunks": 0,
            "total_chunks": 0,
            "created_at": datetime.utcnow().isoformat(),
            "error": None
        })

        return job_id
    
    async def start_bulk_upload(
        self,
        content_type: str,
        items: List[Any],
        brand_id: str
    ) -> str:
        """Start a bulk JSON upload job."""
        job_id = str(uuid.uuid4())
        
        await self.job_store.set(job_id, {
            "status": "pending",
            "type": "bulk",
            "content_type": content_type,
            "brand_id": brand_id,
            "total_items": len(items),
            "processed_items": 0,
            "processed_chunks": 0,
            "created_at": datetime.utcnow().isoformat(),
            "error": None
        })

        return job_id
    
    async def get_job_status(self, job_id: str) -> Optional[Dict]:
        """Get job status."""
        job = await self.job_store.get(job_id)
        if not job:
            return None

        return {
            "job_id": job_id,
            "status": job["status"],
            "progress": {
                "type": job.get("type"),
                "processed_items": job.get("processed_items", 0),
                "total_items": job.get("total_items", 0),
                "processed_chunks": job.get("processed_chunks", 0),
                "total_chunks": job.get("total_chunks", 0)
            },
            "error": job.get("error")
        }
    
    # ========================================================================
    # Document Upload Processing
    # ========================================================================
    
    async def process_document_upload(
        self,
        job_id: str,
        content: bytes,
        filename: str,
        content_type_header: str,
        kb_content_type: str,
        brand_id: str,
        agent_id: Optional[str] = None,
        product_data: Optional[Any] = None,
        dealer_data: Optional[Any] = None,
        folder_path: Optional[str] = None,
    ):
        """Process a document upload in background."""
        try:
            await self._ensure_connection(brand_id)  # Use brand-specific database
            brand_scope = await self._resolve_brand_scope(brand_id)
            configured_default_currency = await self._resolve_configured_default_currency(brand_id, brand_scope)
            await self.job_store.update(job_id, {"status": "processing"})

            # Extract text from document
            text = await self._extract_text(content, content_type_header, filename)
            source_type = self.detect_source_type(content_type_header, filename) or "unknown"

            # Chunk the text (per-agent chunking config applies when agent_id is set)
            chunks = await self._chunk_text(text, filename, agent_id=agent_id)
            await self.job_store.update(job_id, {"total_chunks": len(chunks)})
            
            # Generate doc_id
            doc_id = f"{brand_id}_{kb_content_type}_{uuid.uuid4().hex[:8]}"
            folder, name, path = self._normalize_item_path(folder_path, filename)
            
            # Process each chunk
            for i, chunk_text in enumerate(chunks):
                # Generate embedding
                embedding = await self._generate_embeddings([chunk_text])
                
                # Build chunk document
                chunk_doc = {
                    "doc_id": doc_id,
                    "chunk_id": f"{doc_id}_chunk_{i}",
                    "content": chunk_text,
                    "embeddings": embedding[0] if embedding else [],
                    "title": filename,
                    
                    # Enhanced metadata
                    "content_type": kb_content_type,  # product, dealer, faq, etc.
                    "metadata": {
                        "brand_id": brand_scope.get("brand_id") or brand_id,
                        "brand_slug": brand_scope.get("brand_slug"),
                        "agent_id": agent_id,
                        "filename": filename,
                        "name": name,
                        "folder": folder,
                        "path": path,
                        "content_type_header": content_type_header,
                        "source_type": source_type,
                        "job_id": job_id,
                        "chunk_index": i,
                        "total_chunks": len(chunks),
                        "created_at": datetime.utcnow().isoformat()
                    }
                }
                
                # Add structured data if provided
                if product_data:
                    currency, currency_source = self._resolve_item_currency(product_data, configured_default_currency)
                    chunk_doc["product_data"] = {
                        "sku": product_data.sku,
                        "name": product_data.name,
                        "price": product_data.price,
                        "currency": currency,
                        "currency_source": currency_source,
                        "category": product_data.category,
                        "image_url": product_data.image_url,
                        "product_url": product_data.product_url,
                        "in_stock": product_data.in_stock,
                        "features": product_data.features or []
                    }
                    chunk_doc["product_data"].update(self._product_variant_metadata(product_data))
                
                if dealer_data:
                    chunk_doc["dealer_data"] = {
                        "dealer_id": dealer_data.dealer_id,
                        "name": dealer_data.name,
                        "city": dealer_data.city,
                        "phone": dealer_data.phone,
                        "state": dealer_data.state,
                        "email": dealer_data.email,
                        "address": dealer_data.address
                    }
                
                # Store in MongoDB
                await self.collection.insert_one(chunk_doc)
                if self.qdrant:
                    await self.qdrant.upsert_chunk(chunk_doc, brand_scope.get("brand_slug") or brand_id)
                
                # Update progress
                await self.job_store.update(job_id, {"processed_chunks": i + 1})

            await self.job_store.update(job_id, {"status": "completed"})
            logger.info(
                "Document upload completed",
                job_id=job_id,
                doc_id=doc_id,
                chunks=len(chunks)
            )

        except Exception as e:
            await self.job_store.update(job_id, {"status": "error", "error": str(e)})
            logger.error("Document upload failed", job_id=job_id, error=str(e))
    
    # ========================================================================
    # Bulk Upload Processing
    # ========================================================================
    
    async def process_bulk_upload(
        self,
        job_id: str,
        content_type: str,
        items: List[Any],
        brand_id: str
    ):
        """Process bulk JSON upload in background."""
        try:
            await self._ensure_connection(brand_id)  # Use brand-specific database
            brand_scope = await self._resolve_brand_scope(brand_id)
            configured_default_currency = await self._resolve_configured_default_currency(brand_id, brand_scope)
            await self.job_store.update(job_id, {"status": "processing"})

            for i, item in enumerate(items):
                if content_type == "product":
                    await self._process_product_item(
                        item,
                        brand_id,
                        job_id,
                        brand_scope,
                        configured_default_currency=configured_default_currency,
                    )
                elif content_type == "dealer":
                    await self._process_dealer_item(item, brand_id, job_id, brand_scope)

                # Update progress
                await self.job_store.update(job_id, {"processed_items": i + 1})

            await self.job_store.update(job_id, {"status": "completed"})
            logger.info(
                "Bulk upload completed",
                job_id=job_id,
                items=len(items),
                content_type=content_type
            )

        except Exception as e:
            await self.job_store.update(job_id, {"status": "error", "error": str(e)})
            logger.error("Bulk upload failed", job_id=job_id, error=str(e))
    
    async def _process_product_item(
        self,
        item: Any,
        brand_id: str,
        job_id: str,
        brand_scope: Optional[Dict[str, Any]] = None,
        configured_default_currency: Optional[str] = None,
    ):
        """Process a single product item."""
        currency, currency_source = self._resolve_item_currency(item, configured_default_currency)
        display_price = self._display_product_price(item.price, currency)
        # Generate text description for embedding
        text_parts = [
            f"Product: {item.name}",
            f"SKU: {item.sku}",
            f"Category: {item.category}",
            f"Price: {display_price}"
        ]
        if getattr(item, "parent_name", None) and getattr(item, "parent_name") != item.name:
            text_parts.append(f"Parent Product: {item.parent_name}")
        if getattr(item, "variant_title", None):
            text_parts.append(f"Variant: {item.variant_title}")
        variant_options = getattr(item, "variant_options", None)
        if isinstance(variant_options, dict) and variant_options:
            text_parts.append(
                "Options: " + ", ".join(f"{key}: {value}" for key, value in variant_options.items())
            )
        
        if item.features:
            text_parts.append(f"Features: {', '.join(item.features)}")
        
        full_text = "\n".join(text_parts)
        
        # Generate chunks (split if text is very long)
        chunks = await self._chunk_text(full_text, item.name)
        
        # Generate doc_id based on SKU
        doc_id = f"{brand_id}_product_{item.sku}"
        
        # Process each chunk
        for i, chunk_text in enumerate(chunks):
            # Generate embedding
            embedding = await self._generate_embeddings([chunk_text])
            
            # Build chunk document
            chunk_doc = {
                "doc_id": doc_id,
                "chunk_id": f"{doc_id}_chunk_{i}",
                "content": chunk_text,
                "embeddings": embedding[0] if embedding else [],
                "title": item.name,
                
                # Enhanced metadata
                "content_type": "product",
                "product_data": {
                    "sku": item.sku,
                    "name": item.name,
                    "price": item.price,
                    "currency": currency,
                    "currency_source": currency_source,
                    "category": item.category,
                    "image_url": item.image_url,
                    "product_url": item.product_url,
                    "in_stock": item.in_stock if item.in_stock is not None else True,
                    "features": item.features or []
                },
                "metadata": {
                    "brand_id": (brand_scope or {}).get("brand_id") or brand_id,
                    "brand_slug": (brand_scope or {}).get("brand_slug"),
                    "job_id": job_id,  # Track which upload batch this belongs to
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "created_at": datetime.utcnow().isoformat()
                }
            }
            chunk_doc["product_data"].update(self._product_variant_metadata(item))
            
            # Upsert to MongoDB (update if SKU exists, insert if new)
            await self.collection.update_one(
                {"doc_id": doc_id, "chunk_id": chunk_doc["chunk_id"]},
                {"$set": chunk_doc},
                upsert=True
            )
            if self.qdrant:
                await self.qdrant.upsert_chunk(chunk_doc, (brand_scope or {}).get("brand_slug") or brand_id)
            
            # Update progress (fire-and-forget increment)
            job = await self.job_store.get(job_id) or {}
            await self.job_store.update(job_id, {"processed_chunks": job.get("processed_chunks", 0) + 1})

    def _display_product_price(self, price: Any, currency: Optional[str]) -> str:
        try:
            numeric_price = float(price)
            display_price = numeric_price / 100 if numeric_price >= 10000 else numeric_price
            if display_price.is_integer():
                amount = f"{int(display_price):,}"
            else:
                amount = f"{display_price:,.2f}"
        except (TypeError, ValueError):
            amount = str(price or "0")
        currency = self._normalize_currency(currency)
        return f"{currency} {amount}" if currency else amount
    
    async def _process_dealer_item(self, item: Any, brand_id: str, job_id: str, brand_scope: Optional[Dict[str, Any]] = None):
        """Process a single dealer item."""
        # Generate text description for embedding
        text_parts = [
            f"Dealer: {item.name}",
            f"Dealer ID: {item.dealer_id}",
            f"Location: {item.city}",
            f"Phone: {item.phone}"
        ]
        
        if item.state:
            text_parts.append(f"State: {item.state}")
        if item.email:
            text_parts.append(f"Email: {item.email}")
        if item.address:
            text_parts.append(f"Address: {item.address}")
        
        full_text = "\n".join(text_parts)
        
        # Generate chunks
        chunks = await self._chunk_text(full_text, item.name)
        
        # Generate doc_id based on dealer_id
        doc_id = f"{brand_id}_dealer_{item.dealer_id}"
        
        # Process each chunk
        for i, chunk_text in enumerate(chunks):
            # Generate embedding
            embedding = await self._generate_embeddings([chunk_text])
            
            # Build chunk document
            chunk_doc = {
                "doc_id": doc_id,
                "chunk_id": f"{doc_id}_chunk_{i}",
                "content": chunk_text,
                "embeddings": embedding[0] if embedding else [],
                "title": item.name,
                
                # Enhanced metadata
                "content_type": "dealer",
                "dealer_data": {
                    "dealer_id": item.dealer_id,
                    "name": item.name,
                    "city": item.city,
                    "phone": item.phone,
                    "state": item.state,
                    "email": item.email,
                    "address": item.address
                },
                "metadata": {
                    "brand_id": (brand_scope or {}).get("brand_id") or brand_id,
                    "brand_slug": (brand_scope or {}).get("brand_slug"),
                    "job_id": job_id,  # Track which upload batch this belongs to
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "created_at": datetime.utcnow().isoformat()
                }
            }
            
            # Upsert to MongoDB
            await self.collection.update_one(
                {"doc_id": doc_id, "chunk_id": chunk_doc["chunk_id"]},
                {"$set": chunk_doc},
                upsert=True
            )
            if self.qdrant:
                await self.qdrant.upsert_chunk(chunk_doc, (brand_scope or {}).get("brand_slug") or brand_id)
            
            # Update progress (fire-and-forget increment)
            job = await self.job_store.get(job_id) or {}
            await self.job_store.update(job_id, {"processed_chunks": job.get("processed_chunks", 0) + 1})
    
    # ========================================================================
    # Document Management
    # ========================================================================
    
    async def list_documents(
        self,
        brand_id: str,
        content_type: Optional[str] = None,
        limit: int = 50,
        skip: int = 0
    ) -> List[Dict]:
        """List documents grouped by upload job (not individual products/chunks)."""
        await self._ensure_connection(brand_id)  # Use brand-specific database
        brand_scope = await self._resolve_brand_scope(brand_id)
        brand_aliases = brand_scope.get("aliases") or [brand_id]
        
        # Build query - check both brand_id and brand_slug for compatibility
        brand_filter = {
            "$or": [
                {"metadata.brand_id": {"$in": brand_aliases}},
                {"metadata.brand_slug": {"$in": brand_aliases}},
            ]
        }
        
        if content_type:
            query = {
                "$and": [
                    brand_filter,
                    {"content_type": content_type}
                ]
            }
        else:
            query = brand_filter
        
        logger.info("list_documents_query", brand_id=brand_id, query=query, content_type=content_type)
        
        # Aggregate to group by job_id (upload batch).
        # For old documents without job_id, group by doc_id instead.
        pipeline = [
            {"$match": query},
            {
                "$group": {
                    "_id": {
                        "$ifNull": [
                            "$metadata.job_id",  # New documents have job_id
                            "$doc_id"            # Old documents fall back to doc_id
                        ]
                    },
                    "content_type": {"$first": "$content_type"},
                    "created_at": {"$first": "$metadata.created_at"},
                    "total_chunks": {"$sum": 1},
                    "item_count": {"$addToSet": "$doc_id"},  # Count unique products/dealers
                    "sample_title": {"$first": "$title"},  # Get first item's title
                    "has_job_id": {"$first": {"$ifNull": ["$metadata.job_id", None]}}
                }
            },
            {
                "$project": {
                    "job_id": "$_id",
                    "content_type": 1,
                    "created_at": 1,
                    "total_chunks": 1,
                    "item_count": {"$size": "$item_count"},  # Number of unique items
                    "has_job_id": 1,
                    "title": {
                        "$cond": {
                            "if": {"$ne": ["$has_job_id", None]},
                            "then": {
                                "$concat": [
                                    {"$toString": {"$size": "$item_count"}},
                                    " ",
                                    "$content_type",
                                    "s uploaded"
                                ]
                            },
                            "else": "$sample_title"  # For old documents, use product/dealer name
                        }
                    }
                }
            },
        ]
        
        documents_by_id: Dict[str, Dict[str, Any]] = {}
        for collection in await self._get_brand_knowledge_collections(brand_id):
            cursor = collection.aggregate(pipeline)
            async for doc in cursor:
                doc_id = doc["job_id"] or "unknown"
                existing = documents_by_id.get(doc_id)
                if existing and existing.get("chunks_count", 0) >= doc.get("total_chunks", 0):
                    continue

                documents_by_id[doc_id] = {
                    "doc_id": doc["job_id"] or "unknown",  # Use job_id as doc_id
                    "title": doc.get("title", "Uploaded items"),
                    "content_type": doc.get("content_type"),
                    "chunks_count": doc.get("total_chunks", 0),
                    "item_count": doc.get("item_count", 0),  # Number of products/dealers in this upload
                    "created_at": doc.get("created_at"),
                    "is_legacy": doc.get("has_job_id") is None,  # Flag for old documents
                }
        
        documents = sorted(
            documents_by_id.values(),
            key=lambda doc: doc.get("created_at") or "",
            reverse=True,
        )
        
        return documents[skip:skip + limit]

    async def list_knowledge_tree(
        self,
        brand_id: str,
        agent_id: Optional[str] = None,
        folder: Optional[str] = None,
    ) -> Dict[str, Any]:
        """Return folder and document nodes with filesystem-like paths."""
        brand_scope = await self._resolve_brand_scope(brand_id)
        brand_aliases = brand_scope.get("aliases") or [brand_id]
        selected_folder = self._normalize_folder_path(folder)
        brand_filter: Dict[str, Any] = {
            "$or": [
                {"metadata.brand_id": {"$in": brand_aliases}},
                {"metadata.brand_slug": {"$in": brand_aliases}},
            ]
        }
        if agent_id:
            brand_filter = {"$and": [brand_filter, {"metadata.agent_id": agent_id}]}

        folder_query: Dict[str, Any] = {"brand_id": brand_scope.get("brand_id") or brand_id}
        if agent_id:
            folder_query["agent_id"] = agent_id

        folders_collection = await self._get_knowledge_folders_collection(brand_id)
        folders = [{
            "id": "/",
            "type": "folder",
            "name": "/",
            "path": "/",
            "parent_path": None,
            "created_at": None,
            "updated_at": None,
        }]
        async for folder_doc in folders_collection.find(folder_query).sort("path", 1):
            folders.append({
                "id": folder_doc.get("id") or folder_doc.get("path"),
                "type": "folder",
                "name": folder_doc.get("name"),
                "path": folder_doc.get("path"),
                "parent_path": folder_doc.get("parent_path") or "/",
                "created_at": folder_doc.get("created_at"),
                "updated_at": folder_doc.get("updated_at"),
            })

        documents_by_id: Dict[str, Dict[str, Any]] = {}
        pipeline = [
            {"$match": brand_filter},
            {
                "$group": {
                    "_id": {"$ifNull": ["$metadata.job_id", "$doc_id"]},
                    "doc_id": {"$first": "$doc_id"},
                    "content_type": {"$first": "$content_type"},
                    "title": {"$first": "$title"},
                    "created_at": {"$first": "$metadata.created_at"},
                    "updated_at": {"$max": "$metadata.updated_at"},
                    "folder": {"$first": "$metadata.folder"},
                    "path": {"$first": "$metadata.path"},
                    "name": {"$first": "$metadata.name"},
                    "source_type": {"$first": "$metadata.source_type"},
                    "content_type_header": {"$first": "$metadata.content_type_header"},
                    "agent_id": {"$first": "$metadata.agent_id"},
                    "chunks_count": {"$sum": 1},
                    "item_ids": {"$addToSet": "$doc_id"},
                }
            },
        ]
        for collection in await self._get_brand_knowledge_collections(brand_id):
            async for doc in collection.aggregate(pipeline):
                item_id = doc.get("_id") or doc.get("doc_id")
                if not item_id:
                    continue
                folder_path = self._normalize_folder_path(doc.get("folder"))
                name = doc.get("name") or doc.get("title") or item_id
                path = doc.get("path") or self._normalize_item_path(folder_path, name)[2]
                documents_by_id[str(item_id)] = {
                    "id": str(item_id),
                    "type": "file",
                    "doc_id": str(item_id),
                    "name": name,
                    "title": doc.get("title") or name,
                    "path": path,
                    "folder": folder_path,
                    "parent_path": folder_path,
                    "content_type": doc.get("content_type"),
                    "source_type": doc.get("source_type"),
                    "content_type_header": doc.get("content_type_header"),
                    "agent_id": doc.get("agent_id"),
                    "chunks_count": doc.get("chunks_count", 0),
                    "item_count": len(doc.get("item_ids") or []),
                    "created_at": doc.get("created_at"),
                    "updated_at": doc.get("updated_at") or doc.get("created_at"),
                    "status": "ready",
                }

        all_folders = {folder["path"] for folder in folders if folder.get("path")}
        for document in documents_by_id.values():
            current = document.get("folder") or "/"
            while current and current not in all_folders:
                parent = self._normalize_folder_path("/".join(current.rstrip("/").split("/")[:-1]) or "/")
                folders.append({
                    "id": current,
                    "type": "folder",
                    "name": current.rstrip("/").split("/")[-1] or "/",
                    "path": current,
                    "parent_path": None if current == "/" else parent,
                    "created_at": None,
                    "updated_at": None,
                    "virtual": True,
                })
                all_folders.add(current)
                if current == "/":
                    break
                current = parent

        folder_nodes: Dict[str, Dict[str, Any]] = {}
        for folder_item in folders:
            path = folder_item.get("path") or "/"
            folder_nodes[path] = {
                "id": folder_item.get("id"),
                "name": folder_item.get("name") or (path.rstrip("/").split("/")[-1] or "Knowledge Base"),
                "path": path,
                "parent_id": folder_item.get("parent_path"),
                "parent_path": folder_item.get("parent_path"),
                "children": [],
                "items": [],
                "documents": [],
            }
        folder_nodes.setdefault("/", {
            "id": None,
            "name": "Knowledge Base",
            "path": "/",
            "parent_id": None,
            "parent_path": None,
            "children": [],
            "items": [],
            "documents": [],
        })

        for path, folder_item in sorted(folder_nodes.items(), key=lambda item: item[0]):
            if path == "/":
                continue
            parent_path = folder_item.get("parent_path") or self._normalize_folder_path("/".join(path.rstrip("/").split("/")[:-1]) or "/")
            parent = folder_nodes.setdefault(parent_path, {
                "id": None if parent_path == "/" else parent_path,
                "name": "Knowledge Base" if parent_path == "/" else parent_path.rstrip("/").split("/")[-1],
                "path": parent_path,
                "parent_id": None,
                "parent_path": None,
                "children": [],
                "items": [],
                "documents": [],
            })
            if not any(child.get("path") == path for child in parent["children"]):
                parent["children"].append(folder_item)

        file_items = []
        for document in documents_by_id.values():
            item = {
                **document,
                "kind": "document",
                "id": document.get("id"),
                "source_doc_id": document.get("doc_id"),
                "parent_id": document.get("parent_path"),
            }
            file_items.append(item)
            folder_nodes.setdefault(document.get("folder") or "/", folder_nodes["/"])["items"].append(item)

        children = [
            item for item in [*folder_nodes.get(selected_folder, {}).get("children", []), *folder_nodes.get(selected_folder, {}).get("items", [])]
        ]

        return {
            "root": folder_nodes["/"],
            "selected_folder": selected_folder,
            "folders": sorted(folders, key=lambda item: item.get("path") or ""),
            "files": sorted(file_items, key=lambda item: item.get("updated_at") or "", reverse=True),
            "items": folder_nodes["/"].get("items", []),
            "children": children,
        }

    async def create_folder(
        self,
        brand_id: str,
        path: str,
        agent_id: Optional[str] = None,
    ) -> Dict[str, Any]:
        folder_path = self._normalize_folder_path(path)
        if folder_path == "/":
            raise ValueError("Root folder already exists")
        parent_path = self._normalize_folder_path("/".join(folder_path.rstrip("/").split("/")[:-1]) or "/")
        now = datetime.utcnow().isoformat()
        brand_scope = await self._resolve_brand_scope(brand_id)
        resolved_brand_id = brand_scope.get("brand_id") or brand_id
        folder_doc = {
            "id": folder_path,
            "brand_id": resolved_brand_id,
            "brand_slug": brand_scope.get("brand_slug"),
            "agent_id": agent_id,
            "name": folder_path.rstrip("/").split("/")[-1],
            "path": folder_path,
            "parent_path": parent_path,
            "created_at": now,
            "updated_at": now,
        }
        collection = await self._get_knowledge_folders_collection(brand_id)
        # `updated_at` must appear in exactly one of $setOnInsert / $set — Mongo
        # rejects the same path in both (WriteError 40). Keep it only in $set so
        # the timestamp refreshes on every upsert; the rest is insert-only.
        insert_only = {key: value for key, value in folder_doc.items() if key != "updated_at"}
        await collection.update_one(
            {"brand_id": resolved_brand_id, "agent_id": agent_id, "path": folder_path},
            {"$setOnInsert": insert_only, "$set": {"updated_at": now}},
            upsert=True,
        )
        return {**folder_doc, "type": "folder"}

    async def _find_folder_doc(
        self,
        brand_id: str,
        path: str,
        agent_id: Optional[str] = None,
    ) -> Optional[Dict[str, Any]]:
        """Return the stored folder document for a normalized path, if it exists."""
        brand_scope = await self._resolve_brand_scope(brand_id)
        folders_collection = await self._get_knowledge_folders_collection(brand_id)
        query = {"brand_id": brand_scope.get("brand_id") or brand_id, "path": self._normalize_folder_path(path)}
        if agent_id:
            query["agent_id"] = agent_id
        return await folders_collection.find_one(query)

    async def _cascade_folder_path_change(
        self,
        brand_id: str,
        old_path: str,
        new_path: str,
        agent_id: Optional[str] = None,
    ) -> Dict[str, int]:
        """Re-point every descendant folder and document when a folder moves/renames.

        A folder at ``old_path`` becomes ``new_path``; everything nested under it
        (subfolders + documents) has its stored path prefix rewritten so the tree
        stays consistent. Idempotent and safe when old_path == new_path.
        """
        old_path = self._normalize_folder_path(old_path)
        new_path = self._normalize_folder_path(new_path)
        # Reparenting TO root (new_path == "/") is valid (e.g. folder delete);
        # only the root itself can't move, and same-path is a no-op.
        if old_path == "/" or old_path == new_path:
            return {"folders": 0, "documents": 0}

        now = datetime.utcnow().isoformat()
        old_prefix = old_path + "/"

        def repoint(value: str) -> str:
            if value == old_path:
                return new_path
            if value.startswith(old_prefix):
                # Normalize to avoid '//sub' when new_path is root.
                return self._normalize_folder_path(new_path + value[len(old_path):])
            return value

        # 1. Descendant folder docs (the folder itself is handled by the caller).
        brand_scope = await self._resolve_brand_scope(brand_id)
        folders_collection = await self._get_knowledge_folders_collection(brand_id)
        folder_query: Dict[str, Any] = {
            "brand_id": brand_scope.get("brand_id") or brand_id,
            "path": {"$regex": f"^{re.escape(old_prefix)}"},
        }
        if agent_id:
            folder_query["agent_id"] = agent_id
        folders_updated = 0
        async for child in folders_collection.find(folder_query):
            child_path = child.get("path") or ""
            updated_path = repoint(child_path)
            updated_parent = repoint(child.get("parent_path") or "/")
            await folders_collection.update_one(
                {"_id": child["_id"]},
                {"$set": {"path": updated_path, "id": updated_path, "parent_path": updated_parent, "updated_at": now}},
            )
            folders_updated += 1

        # 2. Documents stored under the old path (in the folder or any descendant).
        brand_aliases = brand_scope.get("aliases") or [brand_id]
        docs_updated = 0
        # Collect Qdrant chunk re-points keyed by destination so the vector
        # payload's folder/path stays in sync with Mongo (folder-scoped retrieval
        # + citations otherwise go stale on the Qdrant backend).
        qdrant_groups: Dict[tuple, list] = {}
        for collection in await self._get_brand_knowledge_collections(brand_id):
            doc_query: Dict[str, Any] = {
                "$and": [
                    {"$or": [
                        {"metadata.folder": old_path},
                        {"metadata.folder": {"$regex": f"^{re.escape(old_prefix)}"}},
                    ]},
                    {"$or": [
                        {"metadata.brand_id": {"$in": brand_aliases}},
                        {"metadata.brand_slug": {"$in": brand_aliases}},
                    ]},
                ]
            }
            if agent_id:
                doc_query["$and"].append({"metadata.agent_id": agent_id})
            async for doc in collection.find(doc_query):
                meta = doc.get("metadata") or {}
                new_folder = repoint(self._normalize_folder_path(meta.get("folder")))
                name = meta.get("name") or doc.get("title") or doc.get("doc_id")
                _, _, new_doc_path = self._normalize_item_path(new_folder, name)
                await collection.update_one(
                    {"_id": doc["_id"]},
                    {"$set": {"metadata.folder": new_folder, "metadata.path": new_doc_path, "metadata.updated_at": now}},
                )
                docs_updated += 1
                if self.qdrant and doc.get("chunk_id"):
                    qdrant_groups.setdefault((new_folder, new_doc_path), []).append(doc["chunk_id"])

        if self.qdrant and qdrant_groups:
            brand_slug = brand_scope.get("brand_slug") or brand_id
            for (new_folder, new_doc_path), chunk_ids in qdrant_groups.items():
                try:
                    await self.qdrant.reparent_chunks(chunk_ids, new_folder, new_doc_path, brand_slug=brand_slug)
                except Exception as e:
                    logger.error("qdrant_reparent_failed", folder=new_folder, error=str(e))

        return {"folders": folders_updated, "documents": docs_updated}

    async def move_item(
        self,
        brand_id: str,
        item_id: str,
        target_folder: str,
        agent_id: Optional[str] = None,
    ) -> Dict[str, Any]:
        target_folder = self._normalize_folder_path(target_folder)
        brand_scope = await self._resolve_brand_scope(brand_id)
        brand_aliases = brand_scope.get("aliases") or [brand_id]
        updated = 0
        item_name = None

        # Folder move: reparent the folder doc under target_folder, then cascade
        # every descendant folder + document to the new path prefix.
        folder_doc = await self._find_folder_doc(brand_id, item_id, agent_id)
        if folder_doc:
            folder_name = folder_doc.get("name") or self._normalize_folder_path(item_id).rstrip("/").split("/")[-1]
            new_path = self._normalize_item_path(target_folder, folder_name)[2]
            if new_path == self._normalize_folder_path(item_id):
                return {"id": new_path, "type": "folder", "path": new_path, "parent_path": target_folder}
            if new_path == folder_doc.get("path") or new_path.startswith((folder_doc.get("path") or "") + "/"):
                raise ValueError("Cannot move a folder into itself")
            now = datetime.utcnow().isoformat()
            folders_collection = await self._get_knowledge_folders_collection(brand_id)
            await folders_collection.update_one(
                {"_id": folder_doc["_id"]},
                {"$set": {"path": new_path, "id": new_path, "parent_path": target_folder, "updated_at": now}},
            )
            await self._cascade_folder_path_change(brand_id, folder_doc.get("path"), new_path, agent_id)
            return {"id": new_path, "type": "folder", "path": new_path, "parent_path": target_folder}

        for collection in await self._get_brand_knowledge_collections(brand_id):
            query = {
                "$and": [
                    {
                        "$or": [
                            {"metadata.job_id": item_id},
                            {"doc_id": item_id},
                        ]
                    },
                    {
                        "$or": [
                            {"metadata.brand_id": {"$in": brand_aliases}},
                            {"metadata.brand_slug": {"$in": brand_aliases}},
                        ]
                    },
                ]
            }
            if agent_id:
                query["$and"].append({"metadata.agent_id": agent_id})
            first = await collection.find_one(query)
            if first and not item_name:
                item_name = (first.get("metadata") or {}).get("name") or first.get("title") or item_id
            folder, name, path = self._normalize_item_path(target_folder, item_name or item_id)
            result = await collection.update_many(
                query,
                {"$set": {
                    "metadata.folder": folder,
                    "metadata.name": name,
                    "metadata.path": path,
                    "metadata.updated_at": datetime.utcnow().isoformat(),
                }},
            )
            updated += result.modified_count

        if updated == 0:
            return {}
        return {"id": item_id, "folder": target_folder, "path": self._normalize_item_path(target_folder, item_name or item_id)[2]}

    async def rename_item(
        self,
        brand_id: str,
        item_id: str,
        name: str,
        agent_id: Optional[str] = None,
    ) -> Dict[str, Any]:
        clean_name = name.strip().strip("/")
        if not clean_name:
            raise ValueError("Name is required")
        brand_scope = await self._resolve_brand_scope(brand_id)
        brand_aliases = brand_scope.get("aliases") or [brand_id]

        folders_collection = await self._get_knowledge_folders_collection(brand_id)
        folder_query = {"brand_id": brand_scope.get("brand_id") or brand_id, "path": item_id}
        if agent_id:
            folder_query["agent_id"] = agent_id
        folder_doc = await folders_collection.find_one(folder_query)
        if folder_doc:
            parent = folder_doc.get("parent_path") or "/"
            old_path = folder_doc.get("path") or item_id
            new_path = self._normalize_item_path(parent, clean_name)[2]
            if new_path != old_path:
                await folders_collection.update_one(
                    {"_id": folder_doc["_id"]},
                    {"$set": {"name": clean_name, "path": new_path, "id": new_path, "updated_at": datetime.utcnow().isoformat()}},
                )
                # Re-point descendant folders + documents to the renamed path.
                await self._cascade_folder_path_change(brand_id, old_path, new_path, agent_id)
            return {"id": new_path, "type": "folder", "name": clean_name, "path": new_path, "parent_path": parent}

        updated = 0
        next_path = None
        for collection in await self._get_brand_knowledge_collections(brand_id):
            query = {
                "$and": [
                    {
                        "$or": [
                            {"metadata.job_id": item_id},
                            {"doc_id": item_id},
                        ]
                    },
                    {
                        "$or": [
                            {"metadata.brand_id": {"$in": brand_aliases}},
                            {"metadata.brand_slug": {"$in": brand_aliases}},
                        ]
                    },
                ]
            }
            if agent_id:
                query["$and"].append({"metadata.agent_id": agent_id})
            first = await collection.find_one(query)
            folder = self._normalize_folder_path((first or {}).get("metadata", {}).get("folder") if first else "/")
            _, _, path = self._normalize_item_path(folder, clean_name)
            next_path = path
            result = await collection.update_many(
                query,
                {"$set": {
                    "title": clean_name,
                    "metadata.filename": clean_name,
                    "metadata.name": clean_name,
                    "metadata.path": path,
                    "metadata.updated_at": datetime.utcnow().isoformat(),
                }},
            )
            updated += result.modified_count

        if updated == 0:
            return {}
        return {"id": item_id, "type": "file", "name": clean_name, "path": next_path}

    async def retrieve(
        self,
        brand_id: str,
        query: str,
        folder: Optional[str] = None,
        agent_id: Optional[str] = None,
        limit: int = 10,
        score_threshold: float = 0.0,
    ) -> List[Dict[str, Any]]:
        """Lightweight admin retrieval preview over stored chunks."""
        brand_scope = await self._resolve_brand_scope(brand_id)
        brand_aliases = brand_scope.get("aliases") or [brand_id]
        query_text = query.strip()
        if not query_text:
            return []

        filters: List[Dict[str, Any]] = [{
            "$or": [
                {"metadata.brand_id": {"$in": brand_aliases}},
                {"metadata.brand_slug": {"$in": brand_aliases}},
            ]
        }]
        if agent_id:
            filters.append({"metadata.agent_id": agent_id})
        if folder:
            folder_path = self._normalize_folder_path(folder)
            filters.append({"metadata.folder": {"$regex": f"^{re.escape(folder_path)}(/|$)"}})

        mongo_query = {"$and": filters}
        tokens = [token.lower() for token in re.findall(r"\w+", query_text) if len(token) > 2]
        results: List[Dict[str, Any]] = []
        seen_chunks = set()

        for collection in await self._get_brand_knowledge_collections(brand_id):
            cursor = collection.find(mongo_query).limit(500)
            async for doc in cursor:
                chunk_id = doc.get("chunk_id") or str(doc.get("_id"))
                if chunk_id in seen_chunks:
                    continue
                seen_chunks.add(chunk_id)
                content = doc.get("content") or ""
                lower_content = content.lower()
                score = 0.0
                if tokens:
                    score = sum(lower_content.count(token) for token in tokens) / max(len(tokens), 1)
                if query_text.lower() in lower_content:
                    score += 3.0
                normalized_score = min(score / 5.0, 1.0)
                if normalized_score < score_threshold:
                    continue
                results.append({
                    "chunk_id": chunk_id,
                    "doc_id": doc.get("doc_id"),
                    "title": doc.get("title"),
                    "content": content,
                    "score": normalized_score,
                    "content_type": doc.get("content_type"),
                    "metadata": {
                        key: value
                        for key, value in (doc.get("metadata") or {}).items()
                        if key not in {"brand_id", "brand_slug"}
                    },
                })

        results.sort(key=lambda item: item.get("score", 0), reverse=True)
        return results[:limit]

    async def get_document_preview(
        self,
        doc_id: str,
        brand_id: str,
        limit: int = 8,
    ) -> Dict[str, Any]:
        """Fetch source metadata plus representative chunks/items for preview."""
        brand_scope = await self._resolve_brand_scope(brand_id)
        brand_aliases = brand_scope.get("aliases") or [brand_id]
        brand_match = {
            "$or": [
                {"metadata.brand_id": {"$in": brand_aliases}},
                {"metadata.brand_slug": {"$in": brand_aliases}},
            ]
        }
        source_match = {
            "$or": [
                {"metadata.job_id": doc_id},
                {"doc_id": doc_id},
            ]
        }
        query = {"$and": [brand_match, source_match]}

        chunks: List[Dict[str, Any]] = []
        seen_chunk_ids = set()
        item_keys = set()
        total_chunks = 0
        collections = await self._get_brand_knowledge_collections(brand_id)

        for collection in collections:
            total_chunks += await collection.count_documents(query)
            for item_key in await collection.distinct("doc_id", query):
                if item_key:
                    item_keys.add(item_key)

        for collection in collections:
            cursor = collection.find(query).sort("metadata.chunk_index", 1).limit(limit)
            async for doc in cursor:
                chunk_id = doc.get("chunk_id") or str(doc.get("_id"))
                if chunk_id in seen_chunk_ids:
                    continue
                seen_chunk_ids.add(chunk_id)
                chunks.append(doc)
                if len(chunks) >= limit:
                    break
            if len(chunks) >= limit:
                break

        if not chunks:
            return {}

        first = chunks[0]
        samples = []
        for chunk in chunks:
            sample: Dict[str, Any] = {
                "chunk_id": chunk.get("chunk_id"),
                "title": chunk.get("title"),
                "content": chunk.get("content"),
                "content_type": chunk.get("content_type"),
                "metadata": {
                    key: value
                    for key, value in (chunk.get("metadata") or {}).items()
                    if key not in {"brand_id", "brand_slug"}
                },
            }
            if chunk.get("product_data"):
                sample["product_data"] = chunk.get("product_data")
            if chunk.get("dealer_data"):
                sample["dealer_data"] = chunk.get("dealer_data")
            samples.append(sample)

        return {
            "doc_id": doc_id,
            "title": first.get("title") or doc_id,
            "content_type": first.get("content_type"),
            "item_count": len(item_keys) or len(samples),
            "chunks_count": total_chunks or len(seen_chunk_ids),
            "created_at": (first.get("metadata") or {}).get("created_at"),
            "status": "ready",
            "samples": samples,
        }
    
    async def delete_document(self, doc_id: str, brand_id: str) -> bool:
        """Delete a document (or job batch) by ID."""
        await self._ensure_connection(brand_id)  # Use brand-specific database
        brand_scope = await self._resolve_brand_scope(brand_id)
        brand_aliases = brand_scope.get("aliases") or [brand_id]
        
        deleted_count = 0
        for collection in await self._get_brand_knowledge_collections(brand_id):
            # First try deleting by job_id (new documents)
            result = await collection.delete_many({
                "metadata.job_id": doc_id,
                "$or": [
                    {"metadata.brand_id": {"$in": brand_aliases}},
                    {"metadata.brand_slug": {"$in": brand_aliases}},
                ]
            })
            deleted_count += result.deleted_count
        
            # If nothing was deleted, try doc_id (old documents without job_id)
            if result.deleted_count == 0:
                result = await collection.delete_many({
                    "doc_id": doc_id,
                    "$or": [
                        {"metadata.brand_id": {"$in": brand_aliases}},
                        {"metadata.brand_slug": {"$in": brand_aliases}},
                    ]
                })
                deleted_count += result.deleted_count

        # When Qdrant is the vector backend, chunks live there too. Mongo deletion
        # alone leaves the vectors searchable, so the "deleted" file would still
        # surface in retrieval. Purge the matching points as well.
        if self.qdrant:
            try:
                await self.qdrant.delete_by_document(
                    doc_id,
                    brand_slug=brand_scope.get("brand_slug") or brand_id,
                    brand_aliases=brand_aliases,
                )
            except Exception as e:
                logger.error("qdrant_delete_failed", doc_id=doc_id, error=str(e))

        return deleted_count > 0

    async def delete_item(self, item_id: str, brand_id: str, agent_id: Optional[str] = None) -> Dict[str, Any]:
        """Delete a knowledge item — folder or document.

        Folders are non-destructive by default: contained documents are reparented
        to the deleted folder's parent so no embedded knowledge is lost, then the
        folder and its descendant folders are removed. Documents delete their chunks.
        """
        folder_doc = await self._find_folder_doc(brand_id, item_id, agent_id)
        if folder_doc:
            return await self.delete_folder(brand_id, folder_doc, agent_id)
        deleted = await self.delete_document(item_id, brand_id=brand_id)
        return {"deleted": deleted, "type": "file"}

    async def delete_folder(
        self,
        brand_id: str,
        folder_doc: Dict[str, Any],
        agent_id: Optional[str] = None,
    ) -> Dict[str, Any]:
        """Remove a folder and its descendant folders; reparent contained docs to the parent."""
        folder_path = self._normalize_folder_path(folder_doc.get("path"))
        parent_path = self._normalize_folder_path(folder_doc.get("parent_path") or "/")

        brand_scope = await self._resolve_brand_scope(brand_id)
        folders_collection = await self._get_knowledge_folders_collection(brand_id)
        prefix = folder_path + "/"

        # Delete the folder records FIRST (before cascade would rename them out of
        # this path prefix), then reparent the contained documents to the parent.
        delete_query: Dict[str, Any] = {
            "brand_id": brand_scope.get("brand_id") or brand_id,
            "$or": [{"path": folder_path}, {"path": {"$regex": f"^{re.escape(prefix)}"}}],
        }
        if agent_id:
            delete_query["agent_id"] = agent_id
        result = await folders_collection.delete_many(delete_query)

        # Reparent every document under this folder (and descendants) up to the parent.
        reparented = await self._cascade_folder_path_change(brand_id, folder_path, parent_path, agent_id)
        return {
            "deleted": result.deleted_count > 0,
            "type": "folder",
            "deleted_folders": result.deleted_count,
            "reparented_documents": reparented.get("documents", 0),
        }

    # ========================================================================
    # Helper Methods
    # ========================================================================

    def detect_source_type(self, content_type: Optional[str], filename: Optional[str]) -> Optional[str]:
        """Normalize an uploaded file type from MIME header or filename extension."""
        normalized_content_type = (content_type or "").split(";", 1)[0].strip().lower()
        if normalized_content_type in self.MIME_SOURCE_TYPES:
            return self.MIME_SOURCE_TYPES[normalized_content_type]

        normalized_filename = (filename or "").lower()
        for extension, source_type in self.EXTENSION_SOURCE_TYPES.items():
            if normalized_filename.endswith(extension):
                return source_type

        return None
    
    async def _extract_text(self, content: bytes, content_type: str, filename: str) -> str:
        """Extract text from different file types."""
        try:
            source_type = self.detect_source_type(content_type, filename)

            if source_type in {"txt", "md"}:
                return self._decode_text(content)
            
            elif source_type == "html":
                return self._extract_html_text(content)
            
            elif source_type == "json":
                # For JSON files, pretty-print the content
                data = json.loads(self._decode_text(content))
                return json.dumps(data, indent=2)
            
            elif source_type == "csv":
                return self._extract_csv_text(content)

            elif source_type == "pdf":
                return self._extract_pdf_text(content)
            
            elif source_type == "docx":
                return self._extract_docx_text(content)
            
            else:
                return self._decode_text(content)
                
        except Exception as e:
            logger.error("Text extraction failed", filename=filename, error=str(e))
            raise Exception(f"Failed to extract text from {filename}: {str(e)}")

    def _decode_text(self, content: bytes) -> str:
        """Decode uploaded text content with a forgiving UTF-8 default."""
        return content.decode("utf-8-sig", errors="replace")

    def _extract_html_text(self, content: bytes) -> str:
        """Extract readable text from lightweight HTML without adding a parser dependency."""
        text = self._decode_text(content)
        text = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", text)
        text = re.sub(r"(?s)<br\s*/?>", "\n", text)
        text = re.sub(r"(?s)</(p|div|section|article|li|tr|h[1-6])>", "\n", text)
        text = re.sub(r"<[^>]+>", " ", text)
        text = html.unescape(text)
        return re.sub(r"[ \t]+\n", "\n", re.sub(r"\n{3,}", "\n\n", text)).strip()

    def _extract_csv_text(self, content: bytes) -> str:
        """Convert CSV rows into labeled plain text for chunking and retrieval."""
        text = self._decode_text(content)
        reader = csv.reader(io.StringIO(text))
        rows = list(reader)
        if not rows:
            return ""

        headers = [header.strip() for header in rows[0]]
        try:
            has_headers = csv.Sniffer().has_header(text[:2048])
        except csv.Error:
            has_headers = any(headers) and len(rows) > 1
        if len(rows) == 1:
            has_headers = False
        extracted_rows = []

        for index, row in enumerate(rows[1:] if has_headers else rows, start=1):
            if has_headers:
                values = []
                for column_index, value in enumerate(row):
                    header = headers[column_index] if column_index < len(headers) and headers[column_index] else f"Column {column_index + 1}"
                    values.append(f"{header}: {value.strip()}")
                extracted_rows.append(f"Row {index}\n" + "\n".join(values))
            else:
                extracted_rows.append(f"Row {index}: " + ", ".join(value.strip() for value in row))

        return "\n\n".join(extracted_rows)

    def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from a PDF using pypdf."""
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError("PDF extraction requires pypdf to be installed") from exc

        reader = PdfReader(io.BytesIO(content))
        page_text = []
        for page_number, page in enumerate(reader.pages, start=1):
            extracted = page.extract_text() or ""
            extracted = extracted.strip()
            if extracted:
                page_text.append(f"Page {page_number}\n{extracted}")

        return "\n\n".join(page_text)

    def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from a DOCX using python-docx."""
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("DOCX extraction requires python-docx to be installed") from exc

        document = Document(io.BytesIO(content))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

        table_rows = []
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_rows.append(" | ".join(cells))

        return "\n\n".join(paragraphs + table_rows)
    
    async def _chunk_text(self, text: str, title: str = "", agent_id: Optional[str] = None) -> List[str]:
        """Chunk text via the shared chunker, honoring the agent's rag.chunking config."""
        chunk_size, chunk_overlap = await resolve_agent_chunking(agent_id)
        return chunk_text(text, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    
    async def _generate_embeddings(self, texts: List[str]) -> List[List[float]]:
        """Generate embeddings using Voyage AI."""
        voyage_config = await self._get_voyage_runtime_config()
        api_key = voyage_config["api_key"]
        base_url = voyage_config["base_url"].rstrip("/")
        model = voyage_config["model"]
        if not api_key:
            logger.warning("voyage_api_key_not_configured", texts_count=len(texts))
            return [[] for _ in texts]

        try:
            async with httpx.AsyncClient(timeout=30.0) as client:
                response = await client.post(
                    f"{base_url}/embeddings",
                    headers={
                        "Authorization": f"Bearer {api_key}",
                        "Content-Type": "application/json"
                    },
                    json={
                        "input": texts,
                        "model": model
                    }
                )
                
                response.raise_for_status()
                data = response.json()
                
                # Extract embeddings
                embeddings = [item["embedding"] for item in data["data"]]
                return embeddings
                
        except Exception as e:
            logger.error("Embedding generation failed", error=str(e))
            # Return empty embeddings on failure (for development)
            return [[] for _ in texts]
